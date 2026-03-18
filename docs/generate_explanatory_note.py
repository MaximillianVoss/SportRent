from __future__ import annotations

import datetime as dt
import json
import os
import sqlite3
from dataclasses import dataclass
from pathlib import Path

import win32com.client  # type: ignore[import]
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
CONTEXT_PATH = REPO_ROOT / "docs" / "explanatory-note-agent-context.json"
DEFAULT_TEMPLATE_PATH = Path(
    r"C:\Users\FossW\OneDrive\Документы\Настраиваемые шаблоны Office\Шаблон пояснительной записки с макросами - с указанием для ИИ агента.dotm"
)
DEFAULT_DB_PATH = REPO_ROOT / "SportRent.Mobile" / "Resources" / "Raw" / "Database" / "sportRent.db"
TEMPLATE_PATH = Path(os.environ.get("SPORTRENT_NOTE_TEMPLATE_PATH", str(DEFAULT_TEMPLATE_PATH)))
DB_PATH = Path(os.environ.get("SPORTRENT_NOTE_DB_PATH", str(DEFAULT_DB_PATH)))
OUTPUT_DIR = REPO_ROOT / "output" / "doc"
SCREENS_DIR = OUTPUT_DIR / "screens"
TMP_DIR = REPO_ROOT / "tmp" / "docs"

SCREEN_FILES = {
    "login": SCREENS_DIR / "01-login.png",
    "catalog": SCREENS_DIR / "02-catalog.png",
    "details": SCREENS_DIR / "03-equipment-details.png",
    "orders": SCREENS_DIR / "04-orders.png",
    "profile": SCREENS_DIR / "05-profile.png",
}

APPENDIX_LABELS = ("А", "Б", "В", "Г")
CODE_LISTINGS = (
    ("Листинг кода инструмента сборки базы данных", REPO_ROOT / "SportRent.DbTool" / "Program.cs"),
    ("Листинг сервиса авторизации пользователя", REPO_ROOT / "SportRent.Mobile" / "Services" / "AuthenticationService.cs"),
    ("Листинг сервиса каталога спортивного инвентаря", REPO_ROOT / "SportRent.Mobile" / "Services" / "SportRentCatalogService.cs"),
    ("Листинг сервиса оформления и чтения заказов", REPO_ROOT / "SportRent.Mobile" / "Services" / "OrdersService.cs"),
)

PLATFORM_ROWS = [
    ["Цель сборки", "Назначение"],
    ["net9.0-android", "Мобильная версия приложения для Android"],
    ["net9.0-ios", "Мобильная версия приложения для iOS"],
    ["net9.0-maccatalyst", "Настольная сборка для macOS через Mac Catalyst"],
    ["net9.0-windows10.0.19041.0", "Десктопная сборка для Windows"],
]

DB_ENTITY_TABLES = [
    (
        "Описание сущностей БД: роли, статусы и способы оплаты",
        [
            ["Сущность", "Назначение", "Ключевые связи"],
            ["roles", "Хранит перечень ролей пользователей.", "Используется внешним ключом users.idRole."],
            [
                "orderStatuses",
                "Содержит справочник статусов заказа аренды.",
                "Используется внешним ключом rentOrders.idStatus.",
            ],
            [
                "paymentMethods",
                "Определяет доступные способы оплаты заказа.",
                "Используется внешним ключом payments.idPaymentMethod.",
            ],
            [
                "paymentStatuses",
                "Фиксирует текущее состояние платежа.",
                "Используется внешним ключом payments.idStatus.",
            ],
            [
                "equipmentConditions",
                "Задает состояния инвентаря, например новый или б/у.",
                "Используется внешним ключом rentalPointEquipment.idEquipmentCondition.",
            ],
        ],
    ),
    (
        "Описание сущностей БД: классификация каталога",
        [
            ["Сущность", "Назначение", "Ключевые связи"],
            [
                "rentalTypes",
                "Описывает варианты тарификации аренды и длительность периода.",
                "Связана с equipmentRates.idRentalType.",
            ],
            ["categories", "Хранит категории спортивного инвентаря.", "Связана с equipment.idCategory."],
            ["brands", "Содержит перечень брендов.", "Связана с equipment.idBrand."],
            ["equipmentTypes", "Определяет типы инвентаря.", "Связана с equipment.idType."],
            [
                "equipmentSizes",
                "Описывает размерные характеристики инвентаря.",
                "Связана с rentalPointEquipment.idSize.",
            ],
        ],
    ),
    (
        "Описание сущностей БД: пользователи, пункты проката и остатки",
        [
            ["Сущность", "Назначение", "Ключевые связи"],
            [
                "rentalPoints",
                "Хранит сведения о пунктах выдачи и возврата инвентаря.",
                "Связана с rentalPointEquipment и rentOrders.",
            ],
            [
                "users",
                "Содержит учетные записи клиентов и сотрудников.",
                "Связана с roles и rentOrders.",
            ],
            [
                "equipment",
                "Описывает конкретные позиции каталога инвентаря.",
                "Связана с categories, brands, equipmentTypes и equipmentRates.",
            ],
            [
                "equipmentRates",
                "Хранит тарифы и размер залога для позиции инвентаря.",
                "Связана с equipment и rentalTypes.",
            ],
            [
                "rentalPointEquipment",
                "Фиксирует остатки инвентаря по пунктам проката, размеру и состоянию.",
                "Связана с rentalPoints, equipment, equipmentConditions и equipmentSizes.",
            ],
        ],
    ),
    (
        "Описание сущностей БД: изображения, заказы и платежи",
        [
            ["Сущность", "Назначение", "Ключевые связи"],
            ["images", "Хранит сведения о локальных графических ресурсах.", "Связана с equipmentPhotos.idImage."],
            [
                "equipmentPhotos",
                "Реализует связь между инвентарем и изображениями.",
                "Связана с equipment и images.",
            ],
            [
                "rentOrders",
                "Содержит заголовок заказа аренды, сроки и суммарные суммы.",
                "Связана с users, orderStatuses и rentalPoints.",
            ],
            [
                "orderItems",
                "Хранит состав заказа по конкретным остаткам инвентаря.",
                "Связана с rentOrders и rentalPointEquipment.",
            ],
            [
                "payments",
                "Фиксирует платеж по заказу, способ и статус оплаты.",
                "Связана с rentOrders, paymentMethods и paymentStatuses.",
            ],
        ],
    ),
]


@dataclass(frozen=True)
class DocumentVariant:
    slug: str
    topic: str
    app_term_genitive: str
    app_term_nominative: str
    scenario_focus: str
    ui_note: str
    testing_note: str
    limitation_note: str
    output_docx_name: str
    output_pdf_name: str


VARIANTS = (
    DocumentVariant(
        slug="mobile",
        topic="Разработка мобильного приложения для аренды спортивного инвентаря",
        app_term_genitive="мобильного приложения",
        app_term_nominative="мобильное приложение",
        scenario_focus=(
            "В данной версии пояснительной записки основной акцент сделан на мобильном сценарии "
            "использования и кроссплатформенной природе решения .NET MAUI."
        ),
        ui_note=(
            "Поскольку проект использует единую XAML-разметку .NET MAUI, в пояснительной записке "
            "используются скриншоты Windows-сборки, отражающие те же экраны и пользовательские "
            "сценарии, что и мобильные цели Android и iOS."
        ),
        testing_note=(
            "Проект ориентирован на мобильные цели net9.0-android и net9.0-ios, однако для "
            "воспроизводимой локальной проверки, подготовки скриншотов и отладки в репозитории "
            "подтверждена сборка Windows-цели net9.0-windows10.0.19041.0. Общая XAML-разметка, "
            "ViewModel и сервисы позволяют рассматривать данную сборку как эквивалентный "
            "демонстрационный контур мобильной версии."
        ),
        limitation_note=(
            "Практическая демонстрация в репозитории выполнена преимущественно на Windows-сборке, "
            "хотя кодовая база ориентирована на мобильные цели Android и iOS."
        ),
        output_docx_name="Пояснительная записка - мобильное приложение SportRent.docx",
        output_pdf_name="Пояснительная записка - мобильное приложение SportRent.pdf",
    ),
    DocumentVariant(
        slug="desktop",
        topic="Разработка десктопного приложения для аренды спортивного инвентаря",
        app_term_genitive="десктопного приложения",
        app_term_nominative="десктопное приложение",
        scenario_focus=(
            "В данной версии пояснительной записки основной акцент сделан на настольной "
            "Windows-сборке кроссплатформенного проекта SportRent."
        ),
        ui_note=(
            "Для десктопной версии в пояснительной записке используются скриншоты Windows-сборки, "
            "так как именно эта конфигурация была применена для отладки, проверки сценариев и "
            "получения иллюстративного материала."
        ),
        testing_note=(
            "Проверка работоспособности выполнялась на Windows-цели net9.0-windows10.0.19041.0. "
            "Именно эта конфигурация рассматривается как десктопная версия приложения в текущей "
            "пояснительной записке. При этом исходная кодовая база остается общей с мобильными "
            "и другими платформенными целями .NET MAUI."
        ),
        limitation_note="Основной подтвержденный сценарий эксплуатации относится к Windows-сборке MAUI.",
        output_docx_name="Пояснительная записка - десктопное приложение SportRent.docx",
        output_pdf_name="Пояснительная записка - десктопное приложение SportRent.pdf",
    ),
)

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
SCREENS_DIR.mkdir(parents=True, exist_ok=True)
TMP_DIR.mkdir(parents=True, exist_ok=True)


def validate_runtime_inputs() -> None:
    missing_inputs: list[str] = []
    if not TEMPLATE_PATH.exists():
        missing_inputs.append(
            f"Word template not found: {TEMPLATE_PATH}. "
            "Set SPORTRENT_NOTE_TEMPLATE_PATH to a valid .dotm file."
        )
    if not DB_PATH.exists():
        missing_inputs.append(
            f"SQLite database not found: {DB_PATH}. "
            "Generate sportRent.db or set SPORTRENT_NOTE_DB_PATH to an existing file."
        )
    if missing_inputs:
        raise FileNotFoundError("\n".join(missing_inputs))


def ensure_base_docx_from_template(template_path: Path, target_path: Path) -> None:
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    word.AutomationSecurity = 3

    document = None
    try:
        document = word.Documents.Add(str(template_path))
        document.SaveAs(str(target_path), FileFormat=16)
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


def export_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    word.AutomationSecurity = 3

    document = None
    try:
        document = word.Documents.Open(str(docx_path))
        document.Fields.Update()
        for toc in document.TablesOfContents:
            toc.Update()
        document.Save()
        document.ExportAsFixedFormat(str(pdf_path), 17)
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


def load_context() -> dict:
    return json.loads(CONTEXT_PATH.read_text(encoding="utf-8"))


def load_facts() -> dict:
    with sqlite3.connect(DB_PATH) as connection:
        cursor = connection.cursor()
        return {
            "table_count": cursor.execute(
                "SELECT count(*) FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'"
            ).fetchone()[0],
            "equipment_count": cursor.execute("SELECT count(*) FROM equipment").fetchone()[0],
            "category_count": cursor.execute("SELECT count(*) FROM categories").fetchone()[0],
            "rental_point_count": cursor.execute("SELECT count(*) FROM rentalPoints").fetchone()[0],
            "user_count": cursor.execute("SELECT count(*) FROM users").fetchone()[0],
            "order_count": cursor.execute("SELECT count(*) FROM rentOrders").fetchone()[0],
            "order_item_count": cursor.execute("SELECT count(*) FROM orderItems").fetchone()[0],
            "payment_count": cursor.execute("SELECT count(*) FROM payments").fetchone()[0],
            "equipment_titles": [
                row[0] for row in cursor.execute("SELECT title FROM equipment ORDER BY id").fetchall()
            ],
        }


def clear_document(document: Document) -> None:
    body = document.element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def get_style_name(document: Document, preferred: list[str], fallback: str) -> str:
    style_names = {style.name for style in document.styles}
    for name in preferred:
        if name in style_names:
            return name
    return fallback


def apply_run_font(run, *, bold: bool | None = None, size: int | None = None) -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    if rpr.rFonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    paragraph.add_run()._r.append(begin)
    paragraph.add_run()._r.append(instr)
    paragraph.add_run()._r.append(separate)
    if placeholder:
        placeholder_run = paragraph.add_run(placeholder)
        apply_run_font(placeholder_run, size=12)
    paragraph.add_run()._r.append(end)


def clear_story_container(container) -> None:
    if hasattr(container, "is_linked_to_previous"):
        container.is_linked_to_previous = False
    container_element = container._element
    for child in list(container_element):
        container_element.remove(child)


def clear_section_footers(section, body_style: str) -> None:
    for footer in (section.footer, section.first_page_footer):
        clear_story_container(footer)
        paragraph = footer.add_paragraph()
        paragraph.style = body_style
        paragraph.paragraph_format.first_line_indent = None


def set_default_section(section, *, different_first_page: bool) -> None:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)
    section.different_first_page_header_footer = different_first_page


def prepare_footer(section, body_style: str) -> None:
    footer = section.footer
    clear_story_container(footer)
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.style = body_style
    paragraph.paragraph_format.first_line_indent = None
    add_field(paragraph, "PAGE", "1")
    if section.different_first_page_header_footer:
        first_page_footer = section.first_page_footer
        clear_story_container(first_page_footer)
        paragraph = first_page_footer.add_paragraph()
        paragraph.style = body_style
        paragraph.paragraph_format.first_line_indent = None


def make_paragraph(
    document: Document,
    text: str = "",
    *,
    style: str,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_cm: float | None = 1.25,
    bold: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.alignment = alignment
    if first_line_indent_cm is None:
        paragraph.paragraph_format.first_line_indent = None
    else:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    run = paragraph.add_run(text)
    apply_run_font(run, bold=bold)


def make_blank(document: Document, body_style: str, count: int = 1) -> None:
    for _ in range(count):
        make_paragraph(document, "", style=body_style, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_cm=None)


def set_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for child in list(ppr):
        if child.tag == qn("w:outlineLvl"):
            ppr.remove(child)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    ppr.append(outline)


def make_chapter_title(document: Document, title: str, style: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    set_outline_level(paragraph, 0)
    run = paragraph.add_run(title)
    apply_run_font(run, bold=True, size=16)


def make_section_title(document: Document, title: str, style: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    set_outline_level(paragraph, 1)
    run = paragraph.add_run(title)
    apply_run_font(run, bold=True, size=14)


def make_centered(document: Document, text: str, style: str, *, bold: bool = False, size: int = 14) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run(text)
    apply_run_font(run, bold=bold, size=size)


def make_list_item(document: Document, text: str, body_style: str) -> None:
    make_paragraph(
        document,
        text,
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_cm=None,
    )


def style_table(table) -> None:
    try:
        table.style = "Сетка таблицы"
    except KeyError:
        table.style = "Table Grid"

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = None
                for run in paragraph.runs:
                    apply_run_font(run, bold=row_index == 0, size=12)


def make_table_with_caption(
    document: Document,
    caption: str,
    rows: list[list[str]],
    caption_style: str,
    body_style: str,
    index: int,
) -> None:
    number_paragraph = document.add_paragraph()
    number_paragraph.style = caption_style
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_paragraph.paragraph_format.first_line_indent = None
    number_run = number_paragraph.add_run(f"Таблица {index}")
    apply_run_font(number_run, size=12)
    make_centered(document, caption, caption_style, size=12)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.style = body_style
                paragraph.paragraph_format.first_line_indent = None
                for run in paragraph.runs:
                    apply_run_font(run, bold=row_index == 0, size=12)
    style_table(table)
    make_blank(document, body_style)


def make_figure(
    document: Document,
    image_path: Path,
    caption: str,
    body_style: str,
    caption_style: str,
    index: int,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = body_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.2))
    make_centered(document, f"Рис. {index}. {caption}", caption_style, size=12)


def add_toc(document: Document, body_style: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = body_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    add_field(
        paragraph,
        r'TOC \o "1-2" \h \z \u',
        "Оглавление будет обновлено при открытии документа.",
    )


def add_main_section(document: Document, body_style: str) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_default_section(section, different_first_page=False)
    prepare_footer(section, body_style)


def add_appendix_heading(document: Document, body_style: str, appendix_label: str, title: str) -> None:
    appendix_paragraph = document.add_paragraph()
    appendix_paragraph.style = body_style
    appendix_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    appendix_paragraph.paragraph_format.first_line_indent = None
    appendix_run = appendix_paragraph.add_run(f"Приложение {appendix_label}")
    apply_run_font(appendix_run, bold=True)

    make_centered(document, title, body_style, bold=True)
    make_blank(document, body_style)


def make_code_listing(document: Document, source_path: Path, body_style: str) -> None:
    code_text = source_path.read_text(encoding="utf-8").replace("\t", "    ")
    numbered_lines = [
        f"{line_no:04}: {line}" if line else f"{line_no:04}:"
        for line_no, line in enumerate(code_text.splitlines(), start=1)
    ]
    paragraph = document.add_paragraph()
    paragraph.style = body_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run("\n".join(numbered_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    rpr = run._element.get_or_add_rPr()
    if rpr.rFonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rpr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.color.rgb = RGBColor(0, 0, 0)


def verify_screenshots_exist() -> None:
    missing = [path for path in SCREEN_FILES.values() if not path.exists()]
    if missing:
        raise FileNotFoundError(f"Missing screenshots: {missing}")


def get_intro_paragraphs(context: dict, variant: DocumentVariant) -> list[str]:
    return [
        (
            "Актуальность разработки приложений для аренды спортивного инвентаря определяется "
            "ростом спроса на краткосрочное пользование оборудованием без необходимости его покупки. "
            "Для клиента важно быстро найти нужную позицию, проверить доступность в конкретном "
            "пункте проката и получить понятные условия аренды."
        ),
        (
            f"Тема курсовой работы связана с созданием {variant.app_term_genitive} на базе "
            "кроссплатформенной технологии .NET MAUI, объединяющего проектирование базы данных "
            "и реализацию пользовательского интерфейса. В рамках проекта SportRent реализован "
            "автономный демонстрационный сценарий, в котором основные действия выполняются на "
            "базе локальной SQLite без серверной части."
        ),
        (
            f"Целью работы является разработка {variant.app_term_genitive} SportRent, "
            "обеспечивающего авторизацию пользователя, просмотр каталога, получение сведений об "
            "инвентаре, оформление аренды, просмотр истории заказов и личного кабинета."
        ),
        (
            "Работа выполнена на примере учебного сервиса проката спортивного инвентаря SportRent, "
            "моделирующего деятельность пункта проката с локальной базой данных и типовыми "
            "пользовательскими сценариями."
        ),
        variant.scenario_focus,
        "Для достижения поставленной цели были решены следующие задачи:",
    ]


def get_task_items(variant: DocumentVariant) -> list[str]:
    task_four = (
        "4. Создано кроссплатформенное мобильное приложение на платформе .NET MAUI."
        if variant.slug == "mobile"
        else "4. Создана десктопная сборка приложения на платформе .NET MAUI для Windows."
    )
    return [
        "1. Выполнен анализ предметной области аренды спортивного инвентаря.",
        "2. Спроектирована нормализованная структура реляционной базы данных.",
        "3. Реализован инструмент автоматической сборки SQLite из SQL-скриптов.",
        task_four,
        "5. Проверена работоспособность ключевых пользовательских сценариев.",
        "6. Выполнен учебный расчет экономической эффективности проекта.",
    ]


def get_requirements_paragraphs(variant: DocumentVariant) -> list[str]:
    platform_requirement = (
        "Для мобильной версии дополнительно важны компактная компоновка элементов, вертикальная "
        "прокрутка контента и удобство сенсорного взаимодействия на экранах ограниченного размера."
        if variant.slug == "mobile"
        else "Для десктопной версии дополнительно важны работа в оконном режиме, рациональное "
        "использование расширенной рабочей области и удобство взаимодействия с помощью клавиатуры и мыши."
    )
    return [
        (
            "На основании анализа предметной области были сформулированы требования к программному "
            "средству. Пользовательская часть должна обеспечивать быстрый вход в систему, поиск и "
            "фильтрацию каталога, просмотр детальной карточки, оформление аренды и получение "
            "информации о ранее созданных заказах."
        ),
        (
            "Система должна работать в автономном режиме, поскольку проект ориентирован на учебную "
            "демонстрацию без зависимости от удаленной инфраструктуры. Поэтому база данных, "
            "изображения инвентаря и логика доступа к данным размещены локально в составе клиента."
        ),
        (
            "Дополнительным требованием стало обеспечение воспроизводимости данных. Для этого в "
            "решении выделен отдельный инструмент сборки базы данных из канонических SQL-скриптов "
            "и подготовлены демонстрационные учетные записи разных ролей."
        ),
        platform_requirement,
    ]


def get_architecture_paragraphs(variant: DocumentVariant) -> list[str]:
    focus_text = (
        "В рамках данной версии записки основной акцент сделан на мобильных пользовательских "
        "сценариях и переносимости экранов между Android и iOS."
        if variant.slug == "mobile"
        else "В рамках данной версии записки основной акцент сделан на Windows-сборке, которая "
        "рассматривается как десктопный вариант эксплуатации приложения."
    )
    return [
        (
            "Решение состоит из трех логически разделенных проектов. Такое разделение позволяет "
            "обособить задачи построения базы данных, работы со scaffolded сущностями и разработки "
            "пользовательского интерфейса."
        ),
        (
            "В клиентском приложении используется упрощенный вариант архитектурного подхода MVVM: "
            "страницы отвечают за визуальную часть, ViewModel управляют состоянием экрана, а сервисы "
            "инкапсулируют доступ к данным и прикладную логику. Такой подход уменьшает связанность "
            "кода и делает сценарии загрузки каталога, авторизации и заказов более прозрачными."
        ),
        (
            "Файл SportRent.Mobile.csproj подтверждает кроссплатформенный характер решения: "
            "проект ориентирован на Android, iOS, Mac Catalyst и Windows. Общая XAML-разметка и "
            "единый слой сервисов обеспечивают использование одной кодовой базы для разных целей сборки."
        ),
        focus_text,
        (
            "Точкой входа служит класс App, который определяет стартовый экран по состоянию "
            "пользовательской сессии. При отсутствии авторизации отображается LoginPage, а после "
            "входа открывается AppShell с вкладками «Каталог», «Заказы» и «Профиль»."
        ),
    ]


def get_interface_intro_paragraphs(variant: DocumentVariant) -> list[str]:
    layout_note = (
        "Интерфейс мобильной версии ориентирован на вертикальный сценарий просмотра: блоки "
        "последовательно выстраиваются сверху вниз, а ключевые действия размещены в видимой зоне "
        "экрана, что упрощает сенсорное взаимодействие."
        if variant.slug == "mobile"
        else "Интерфейс десктопной версии работает в оконном режиме Windows и использует увеличенную "
        "ширину рабочей области для более комфортного просмотра карточек, списков и сводных показателей."
    )
    return [
        (
            "Интерфейс приложения оформлен в едином визуальном стиле и ориентирован на быстрое "
            "прохождение основных сценариев. На экранах используются крупные карточки, сводные "
            "метрики, фильтрация по категориям и визуально выделенные области для стоимости, "
            "залога и доступности."
        ),
        layout_note,
        variant.ui_note,
    ]


def get_testing_paragraphs(variant: DocumentVariant, facts: dict) -> list[str]:
    return [
        variant.testing_note,
        (
            f"В результате проверки установлено, что приложение корректно отображает "
            f"{facts['equipment_count']} позиций каталога, сгруппированных по "
            f"{facts['category_count']} категориям и доступных в {facts['rental_point_count']} "
            "пунктах проката. История заказов показывает подготовленный тестовый заказ, а профиль "
            "пользователя агрегирует статистику по заказам и платежам."
        ),
        (
            "Основной проверочный сценарий включает вход под демонстрационным пользователем, "
            "просмотр каталога, открытие карточки инвентаря, переход к истории заказов и просмотру "
            "личного кабинета. Этот набор действий показывает связность интерфейса, корректность "
            "навигации и согласованность локальных данных."
        ),
    ]


def get_limitations(variant: DocumentVariant) -> list[str]:
    return [
        "1. Авторизация демонстрационная, пароль хранится в seed-данных без криптографического хеширования.",
        "2. Серверная часть и синхронизация с удаленной базой данных отсутствуют.",
        "3. Регистрация новых пользователей и полноценная ролевая модель интерфейса не реализованы.",
        "4. Реальная интеграция с платежной системой отсутствует.",
        "5. Автоматизированные тесты в текущей версии не разработаны.",
        f"6. {variant.limitation_note}",
    ]


def get_economic_cost_rows() -> list[list[str]]:
    return [
        ["Статья затрат", "Основание расчета", "Сумма, руб."],
        ["Анализ и проектирование", "48 ч x 350 руб./ч", "16 800"],
        ["Разработка базы данных и сервисов", "92 ч x 350 руб./ч", "32 200"],
        ["Разработка пользовательского интерфейса", "76 ч x 350 руб./ч", "26 600"],
        ["Тестирование и отладка", "40 ч x 320 руб./ч", "12 800"],
        ["Электроэнергия и амортизация ПК", "Условный расчет", "4 600"],
        ["Итого", "", "93 000"],
    ]


def get_economic_effect_rows() -> list[list[str]]:
    return [
        ["Показатель", "Значение"],
        ["Среднее сокращение времени на оформление одной аренды", "8 минут"],
        ["Среднее количество заявок в месяц", "180"],
        ["Экономия рабочего времени в месяц", "24 часа"],
        ["Средняя стоимость часа работы сотрудника", "350 руб."],
        ["Условная ежемесячная экономия", "8 400 руб."],
        ["Дополнительный эффект от снижения ошибок и повторных операций", "4 000 руб."],
        ["Совокупный экономический эффект в месяц", "12 400 руб."],
        ["Ориентировочный срок окупаемости", "около 7,5 месяцев"],
    ]


def get_conclusion_paragraphs(variant: DocumentVariant) -> list[str]:
    first = (
        "В ходе выполнения курсовой работы разработано мобильное кроссплатформенное приложение "
        "SportRent для аренды спортивного инвентаря с локальной SQLite-базой данных."
        if variant.slug == "mobile"
        else "В ходе выполнения курсовой работы разработана десктопная версия кроссплатформенного "
        "приложения SportRent для аренды спортивного инвентаря с локальной SQLite-базой данных."
    )
    return [
        first,
        (
            f"В результате работы спроектирована нормализованная база данных, создан инструмент ее "
            "автоматической сборки из SQL-скриптов, выполнено scaffold-моделирование слоя данных и "
            f"реализовано {variant.app_term_nominative} на .NET MAUI. Пользователю доступны "
            "авторизация, каталог, карточка инвентаря, оформление аренды, история заказов и профиль."
        ),
        (
            f"Поставленная цель достигнута, так как тема курсовой работы «{variant.topic}» "
            "подтверждена фактической реализацией программного средства. Полученный результат может "
            "служить основой для дальнейшего развития проекта, включая добавление серверной части, "
            "улучшение безопасности, расширение ролей пользователей и интеграцию с внешними сервисами."
        ),
    ]


def build_document(variant: DocumentVariant) -> tuple[Path, Path]:
    context = load_context()
    facts = load_facts()
    verify_screenshots_exist()

    base_docx_path = TMP_DIR / f"{variant.slug}-template-base.docx"
    docx_path = OUTPUT_DIR / variant.output_docx_name
    pdf_path = OUTPUT_DIR / variant.output_pdf_name

    ensure_base_docx_from_template(TEMPLATE_PATH, base_docx_path)
    document = Document(str(base_docx_path))
    clear_document(document)

    body_style = get_style_name(document, ["ГОСТ Текст", "Основной текст", "Normal"], "Normal")
    chapter_style = body_style
    section_style = body_style
    caption_style = get_style_name(document, ["caption", "Caption", "ГОСТ Текст"], body_style)

    for section in document.sections:
        set_default_section(section, different_first_page=True)
        clear_section_footers(section, body_style)

    year = str(dt.date.today().year)
    figure_index = 1
    table_index = 1

    make_centered(document, "[Наименование образовательной организации]", body_style, bold=True)
    make_centered(document, "[Кафедра / отделение]", body_style)
    make_blank(document, body_style, 4)
    make_centered(document, "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", body_style, bold=True, size=18)
    make_centered(document, "к курсовой работе", body_style)
    make_centered(document, "на тему", body_style)
    make_centered(document, f"«{variant.topic}»", body_style, bold=True)
    make_blank(document, body_style, 3)
    make_paragraph(
        document,
        f"Проект: {context['project_name']}",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent_cm=None,
    )
    make_paragraph(
        document,
        f"Выполнил: {context['author']['full_name']}",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent_cm=None,
    )
    make_paragraph(
        document,
        "Группа: ____________________",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent_cm=None,
    )
    make_paragraph(
        document,
        "Руководитель: ____________________",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent_cm=None,
    )
    make_blank(document, body_style, 3)
    make_centered(document, year, body_style)

    document.add_page_break()
    make_chapter_title(document, "СОДЕРЖАНИЕ", chapter_style)
    add_toc(document, body_style)

    add_main_section(document, body_style)
    make_chapter_title(document, "ВВЕДЕНИЕ", chapter_style)
    for paragraph in get_intro_paragraphs(context, variant):
        make_paragraph(document, paragraph, style=body_style)
    for item in get_task_items(variant):
        make_list_item(document, item, body_style)
    make_paragraph(
        document,
        (
            "Объектом исследования является процесс аренды спортивного инвентаря, а предметом "
            "исследования выступают методы проектирования локальной базы данных и реализации "
            f"{variant.app_term_genitive} для обслуживания данного процесса."
        ),
        style=body_style,
    )
    make_paragraph(
        document,
        "В работе использованы методы анализа предметной области, сравнения аналогичных решений, проектирования структуры данных, программной реализации, тестирования и условного экономического расчета.",
        style=body_style,
    )

    document.add_page_break()
    make_chapter_title(document, "ГЛАВА 1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ ИЗУЧАЕМОЙ ПРОБЛЕМЫ", chapter_style)
    make_section_title(document, "1.1 Характеристика предметной области", section_style)
    chapter_one = [
        "Предметная область включает пользователей, пункты проката, категории и конкретные позиции инвентаря, тарифы аренды, остатки, заказы и платежи. Для одного и того же вида инвентаря необходимо хранить сведения о доступности в разных пунктах, размерах и состояниях, а также учитывать залог и срок аренды.",
        f"В текущей версии проекта в локальной базе данных хранится {facts['category_count']} категорий, {facts['equipment_count']} позиций инвентаря и {facts['rental_point_count']} пункта проката. Пользователю доступны такие примеры оборудования, как {', '.join(facts['equipment_titles'][:5])}. Эти данные достаточны для демонстрации типового потока работы приложения на защите.",
        "Ключевая особенность рассматриваемой предметной области состоит в необходимости связать каталог с реальными остатками инвентаря. Пользователь должен видеть не только карточку товара, но и доступность выбранной позиции в конкретном пункте проката, а также применимые тарифы и размер залога.",
    ]
    for paragraph in chapter_one:
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "1.2 Требования к цифровому сервису аренды", section_style)
    for paragraph in get_requirements_paragraphs(variant):
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "1.3 Обоснование выбора средств разработки", section_style)
    for paragraph in [
        f"Для реализации проекта выбран стек, включающий {context['technology_stack']['platform']} в качестве платформы, {context['technology_stack']['ui']} для построения интерфейса, язык {context['technology_stack']['language']}, разметку {context['technology_stack']['markup']}, СУБД {context['technology_stack']['database']} и Entity Framework Core 9 для слоя данных. Такой набор средств соответствует теме курсовой работы и позволяет объединить разработку интерфейса, локального хранения и прикладной логики в одной технологии.",
        "Выбор SQLite обусловлен малым объемом демонстрационных данных, простотой развертывания и возможностью хранить файл базы непосредственно в составе приложения. .NET MAUI обеспечивает кроссплатформенную основу и единый стек для реализации страниц, навигации и взаимодействия с локальными ресурсами.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Использованные технологии и их назначение",
        [
            ["Технология", "Назначение"],
            [".NET 9", "Платформа разработки и выполнения проекта"],
            [".NET MAUI", "Построение кроссплатформенного пользовательского интерфейса"],
            ["C#", "Реализация прикладной логики и сервисов"],
            ["XAML", "Описание экранов приложения"],
            ["SQLite", "Локальное хранение данных"],
            ["Entity Framework Core 9", "Scaffolded слой данных и модели"],
            ["PowerShell", "Автоматизация сборки базы данных и scaffold-процесса"],
        ],
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1

    document.add_page_break()
    make_chapter_title(document, "ГЛАВА 2. ПРОЕКТИРОВАНИЕ СТРУКТУРЫ ПРИЛОЖЕНИЯ", chapter_style)
    make_section_title(document, "2.1 Проектирование базы данных", section_style)
    for paragraph in [
        f"База данных проекта сформирована как нормализованная реляционная структура и в текущем состоянии содержит {facts['table_count']} таблиц. При проектировании отдельно выделены справочники ролей, категорий, брендов, типов аренды, статусов заказов и платежей, что исключает дублирование данных и упрощает дальнейшее расширение системы.",
        "Основной сущностью каталога является таблица equipment, содержащая сведения о конкретной позиции инвентаря. Тарифы вынесены в отдельную таблицу equipmentRates, что позволяет задавать несколько вариантов аренды для одной позиции. Остатки хранятся в rentalPointEquipment и связаны одновременно с пунктом проката, размером и состоянием оборудования.",
        "Для корректного хранения медиафайлов реализована нормализация изображений через таблицы images и equipmentPhotos. Позиции заказа связаны не просто с инвентарем вообще, а с конкретным остатком из rentalPointEquipment, благодаря чему приложение может корректно уменьшать доступное количество при создании заказа.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Ключевые группы таблиц базы данных",
        [
            ["Группа", "Содержание"],
            ["Пользователи и роли", "roles, users"],
            ["Каталог и классификаторы", "categories, brands, equipmentTypes, equipmentSizes, equipmentConditions"],
            ["Прокат и доступность", "rentalTypes, rentalPoints, equipment, equipmentRates, rentalPointEquipment"],
            ["Изображения", "images, equipmentPhotos"],
            ["Заказы и статусы", "orderStatuses, rentOrders, orderItems"],
            ["Платежи", "paymentMethods, paymentStatuses, payments"],
        ],
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    make_paragraph(
        document,
        "Важным проектным решением стало хранение денежных значений в целочисленном формате, то есть в копейках. Это исключает ошибки округления и делает операции расчета аренды и залога более предсказуемыми.",
        style=body_style,
    )
    make_paragraph(
        document,
        "Ниже приведено краткое описание основных сущностей базы данных и их связей.",
        style=body_style,
    )
    for caption, rows in DB_ENTITY_TABLES:
        make_table_with_caption(document, caption, rows, caption_style, body_style, table_index)
        table_index += 1
    make_paragraph(
        document,
        "Представленные сущности покрывают как справочный контур системы, так и операционный контур аренды, включая остатки, заказы и платежи. Благодаря этому схема БД соответствует предметной области и поддерживает полный учебный сценарий.",
        style=body_style,
    )
    make_section_title(document, "2.2 Проектирование архитектуры решения", section_style)
    for paragraph in get_architecture_paragraphs(variant):
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Состав решения SportRent",
        [
            ["Проект", "Назначение"],
            ["SportRent.DbTool", "Сборка SQLite-базы из упорядоченных SQL-скриптов"],
            ["SportRent.Data", "Scaffolded EF Core сущности и DbContext"],
            ["SportRent.Mobile", "Кроссплатформенное приложение с экранами, сервисами и навигацией для мобильной и десктопной сборки"],
        ],
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    make_table_with_caption(
        document,
        "Целевые платформы единой кодовой базы MAUI",
        PLATFORM_ROWS,
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    document.add_page_break()
    make_chapter_title(document, "ГЛАВА 3. РЕАЛИЗАЦИЯ И ТЕСТИРОВАНИЕ ПРИЛОЖЕНИЯ", chapter_style)
    make_section_title(document, "3.1 Реализация основных пользовательских сценариев", section_style)
    for paragraph in [
        "Сервис авторизации обращается к локальной базе данных, загружает демонстрационные учетные записи и выполняет проверку логина и пароля по таблице users с учетом роли пользователя. В учебной версии проекта учетные данные используются в демонстрационном виде, что упрощает проверку интерфейса и базы данных.",
        "Сервис каталога формирует снимок данных для главного экрана, включающий список категорий, оборудование и сводные показатели по количеству позиций, пунктов проката и категорий. Для карточки инвентаря сервис дополнительно извлекает тарифы аренды, величину залога и сведения о доступности по точкам проката.",
        "Наиболее важным с точки зрения прикладной логики является сценарий оформления аренды. В сервисе заказов реализована транзакционная последовательность действий: проверка остатка, уменьшение доступного количества, создание записи заказа, добавление позиции заказа и автоматическое создание платежа со статусом ожидания оплаты. Такой порядок обеспечивает согласованность локальных данных.",
        f"Подготовленные демонстрационные данные содержат {facts['user_count']} учетные записи, {facts['order_count']} тестовый заказ, {facts['order_item_count']} позицию(й) заказа и {facts['payment_count']} платежную запись. Благодаря этому пользователь может показать на защите не только каталог, но и уже заполненную историю аренды.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "3.2 Разработка интерфейса приложения", section_style)
    for paragraph in get_interface_intro_paragraphs(variant):
        make_paragraph(document, paragraph, style=body_style)
    make_paragraph(
        document,
        "Экран авторизации.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "Первым для пользователя открывается экран входа, на котором размещены предзаполненные поля почты и пароля, а также список демонстрационных учетных записей. Такое решение уменьшает время подготовки к защите и позволяет быстро перейти к показу основных функций приложения.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["login"],
        "Экран входа в систему с демонстрационными учетными записями",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Главный экран каталога.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "После успешной авторизации открывается вкладка каталога. На ней отображаются агрегированные показатели по количеству позиций, пунктов проката и категорий, строка поиска, горизонтальный набор фильтров и карточки инвентаря с кратким описанием, стартовой ценой и залогом.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["catalog"],
        "Главный экран каталога спортивного инвентаря",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Карточка инвентаря и переход к оформлению аренды.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "При выборе позиции из каталога пользователь переходит на экран подробной карточки. Здесь собраны сведения о категории, бренде и модели, минимальной стоимости аренды, размере залога, доступности по пунктам проката и доступных тарифах. С этого экрана выполняется переход к отдельной форме оформления аренды, где пользователь выбирает тариф, пункт выдачи, дату начала, количество единиц и сразу видит итоговый расчет платежа.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["details"],
        "Карточка выбранной позиции инвентаря",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Экран истории заказов.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "Раздел заказов показывает уже оформленные аренды и предназначен для проверки результатов транзакционного сценария. На экране видны статус заказа, сроки аренды, сумма аренды, размер залога, итоговый платеж и состояние оплаты, что позволяет наглядно проверить корректность записи данных в локальную базу.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["orders"],
        "Экран истории заказов пользователя",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Экран профиля пользователя.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "Вкладка профиля завершает пользовательский маршрут и содержит персональные данные, роль, контакты, дату регистрации и сводные показатели по заказам и оплатам. Наличие кнопки выхода из аккаунта подтверждает завершенность базового сценария работы с пользовательской сессией.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["profile"],
        "Экран профиля пользователя",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Представленные экраны показывают, что тема курсовой работы соответствует фактически реализованному приложению: пользователь может войти в систему, выбрать инвентарь, ознакомиться с условиями аренды, посмотреть историю заказов и получить сведения о своем профиле.",
        style=body_style,
    )

    document.add_page_break()
    make_section_title(document, "3.3 Тестирование и проверка работоспособности", section_style)
    make_section_title(document, "3.3.1 Сборка локальной базы данных", section_style)
    for paragraph in [
        "Для подготовки базы данных реализован отдельный консольный инструмент SportRent.DbTool. Он выполняет поиск SQL-файлов формата 000_name.sql в каталоге database, сортирует их по имени, создает временный файл SQLite и последовательно выполняет все найденные сценарии в рамках транзакции.",
        "В проекте используются три канонических сценария: 001_recreate.sql, 002_schema.sql и 003_seed.sql. Первый сценарий отвечает за пересоздание структуры, второй - за формирование схемы, третий - за наполнение демонстрационными данными.",
        "После успешного выполнения SQL-скриптов временный файл базы данных заменяет итоговый файл sportRent.db в ресурсах мобильного приложения. Данный подход снижает риск повреждения рабочей базы и делает процесс подготовки данных воспроизводимым.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "3.3.2 Проверка пользовательских сценариев", section_style)
    for paragraph in get_testing_paragraphs(variant, facts):
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "3.3.3 Ограничения текущей версии", section_style)
    make_paragraph(
        document,
        "Рассматриваемая реализация является учебной и не претендует на уровень промышленной системы. Это следует учитывать при оценке результатов и определении направлений дальнейшего развития проекта.",
        style=body_style,
    )
    for item in get_limitations(variant):
        make_list_item(document, item, body_style)
    make_paragraph(
        document,
        "Несмотря на перечисленные ограничения, проект решает поставленные учебные задачи и демонстрирует полный маршрут обработки данных от структуры базы до экранов конечного пользователя.",
        style=body_style,
    )

    document.add_page_break()
    make_chapter_title(document, "ГЛАВА 4. РАСЧЕТ ЭКОНОМИЧЕСКОЙ ЭФФЕКТИВНОСТИ ПРОЕКТА", chapter_style)
    make_section_title(document, "4.1 Анализ затрат на разработку", section_style)
    for paragraph in [
        "Расчет экономической эффективности выполнен в учебных целях и носит оценочный характер. Он позволяет показать, какие трудозатраты и сопутствующие расходы могут сопровождать разработку подобного программного продукта.",
        "В расчет включены основные этапы проектирования, программной реализации, тестирования, а также условные затраты на использование вычислительной техники. Оценка построена на трудоемкости работ, наблюдаемой в ходе разработки учебного проекта SportRent.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Оценка затрат на разработку проекта",
        get_economic_cost_rows(),
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    make_section_title(document, "4.2 Оценка экономической эффективности", section_style)
    for paragraph in [
        "Для оценки эффекта рассмотрен типовой пункт проката, в котором часть операций менеджера переносится в приложение: просмотр каталога, уточнение доступности, расчет стоимости аренды и фиксация заказа. Это уменьшает время обслуживания одного клиента и снижает вероятность повторного ввода данных.",
        "Показатели экономического эффекта основаны на условных, но реалистичных допущениях по количеству обращений и стоимости рабочего времени сотрудника. Такой подход соответствует учебному характеру проекта и позволяет определить ориентировочный срок окупаемости.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Оценка ежемесячного экономического эффекта",
        get_economic_effect_rows(),
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    make_paragraph(
        document,
        "Полученные оценки показывают, что даже при учебном масштабе проекта приложение способно сократить операционные затраты пункта проката и окупить условные затраты на разработку в разумный срок. Следовательно, разработка SportRent является не только технически реализуемой, но и экономически целесообразной в модели демонстрационного внедрения.",
        style=body_style,
    )

    document.add_page_break()
    make_chapter_title(document, "ЗАКЛЮЧЕНИЕ", chapter_style)
    for paragraph in get_conclusion_paragraphs(variant):
        make_paragraph(document, paragraph, style=body_style)

    document.add_page_break()
    make_chapter_title(document, "СПИСОК ЛИТЕРАТУРЫ", chapter_style)
    for item in [
        "1. Методические указания к дипломному проектированию. Локальный документ «МУ к ДП.pdf».",
        "2. Репозиторий SportRent. Файл README.md.",
        "3. Репозиторий SportRent. Файл docs/explanatory-note-agent-context.json.",
        "4. Репозиторий SportRent. Файл SportRent.Mobile/SportRent.Mobile.csproj.",
        "5. Репозиторий SportRent. Исходный код проекта SportRent.DbTool, файл Program.cs.",
        "6. Репозиторий SportRent. Скрипт SportRent.Data/scaffold.ps1.",
        "7. Репозиторий SportRent. Исходный код сервиса авторизации SportRent.Mobile/Services/AuthenticationService.cs.",
        "8. Репозиторий SportRent. Исходный код каталожного сервиса SportRent.Mobile/Services/SportRentCatalogService.cs.",
        "9. Репозиторий SportRent. Исходный код сервиса заказов SportRent.Mobile/Services/OrdersService.cs.",
        "10. Репозиторий SportRent. SQL-сценарий структуры базы данных database/002_schema.sql.",
        "11. Репозиторий SportRent. SQL-сценарий начального заполнения database/003_seed.sql.",
        "12. Microsoft Learn. Documentation for .NET MAUI.",
        "13. Microsoft Learn. Documentation for Entity Framework Core.",
        "14. SQLite Documentation.",
    ]:
        make_list_item(document, item, body_style)

    for appendix_label, (appendix_title, appendix_path) in zip(APPENDIX_LABELS, CODE_LISTINGS, strict=True):
        document.add_page_break()
        add_appendix_heading(
            document,
            body_style,
            appendix_label,
            f"{appendix_title} ({appendix_path.name})",
        )
        make_code_listing(document, appendix_path, body_style)

    document.save(docx_path)
    export_to_pdf(docx_path, pdf_path)
    return docx_path, pdf_path


if __name__ == "__main__":
    validate_runtime_inputs()
    for current_variant in VARIANTS:
        current_docx, current_pdf = build_document(current_variant)
        print(f"{current_variant.slug.upper()} DOCX: {current_docx}")
        print(f"{current_variant.slug.upper()} PDF: {current_pdf}")
