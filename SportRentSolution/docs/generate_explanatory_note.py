from __future__ import annotations

import datetime as dt
import json
import os
import sqlite3
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import win32com.client  # type: ignore[import]
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[1]
CONTEXT_PATH = REPO_ROOT / "docs" / "explanatory-note-agent-context.json"
DEFAULT_TEMPLATE_PATH = Path(
    r"C:\Users\FossW\OneDrive\Документы\Настраиваемые шаблоны Office\Шаблон пояснительной записки с макросами - с указанием для ИИ агента.dotm"
)
DEFAULT_DB_PATH = REPO_ROOT / "SportRent.Mobile" / "Resources" / "Raw" / "Database" / "sportRent.db"
TEMPLATE_PATH = Path(os.environ.get("SPORTRENT_NOTE_TEMPLATE_PATH", str(DEFAULT_TEMPLATE_PATH)))
DB_PATH = Path(os.environ.get("SPORTRENT_NOTE_DB_PATH", str(DEFAULT_DB_PATH)))
OUTPUT_DIR = REPO_ROOT / "output" / "doc"
SCREENS_DIR = OUTPUT_DIR / "screens"
TMP_DIR = REPO_ROOT / "tmp" / "docs"

SCREEN_FILES = {
    "login": SCREENS_DIR / "01-login.png",
    "catalog": SCREENS_DIR / "02-catalog.png",
    "details": SCREENS_DIR / "03-equipment-details.png",
    "orders": SCREENS_DIR / "04-orders.png",
    "profile": SCREENS_DIR / "05-profile.png",
}

APPENDIX_LABELS = ("А", "Б", "В", "Г")
CODE_LISTINGS = (
    ("Листинг кода инструмента сборки базы данных", REPO_ROOT / "SportRent.DbTool" / "Program.cs"),
    ("Листинг сервиса авторизации пользователя", REPO_ROOT / "SportRent.Mobile" / "Services" / "AuthenticationService.cs"),
    ("Листинг сервиса каталога спортивного инвентаря", REPO_ROOT / "SportRent.Mobile" / "Services" / "SportRentCatalogService.cs"),
    ("Листинг сервиса оформления и чтения заказов", REPO_ROOT / "SportRent.Mobile" / "Services" / "OrdersService.cs"),
)

PLATFORM_ROWS = [
    ["Цель сборки", "Назначение"],
    ["net9.0-android", "Мобильная версия приложения для Android"],
    ["net9.0-ios", "Мобильная версия приложения для iOS"],
    ["net9.0-maccatalyst", "Настольная сборка для macOS через Mac Catalyst"],
    ["net9.0-windows10.0.19041.0", "Десктопная сборка для Windows"],
]

DB_TABLE_DESCRIPTION_MAP = {
    "roles": "справочник ролей пользователей",
    "orderStatuses": "справочник статусов заказов аренды",
    "paymentMethods": "справочник способов оплаты",
    "paymentStatuses": "справочник статусов платежей",
    "equipmentConditions": "справочник состояний инвентаря",
    "rentalTypes": "варианты тарификации аренды и длительность периода",
    "categories": "категории спортивного инвентаря",
    "brands": "справочник брендов оборудования",
    "equipmentTypes": "классификатор типов инвентаря",
    "equipmentSizes": "классификатор размеров инвентаря",
    "rentalPoints": "пункты выдачи и возврата инвентаря",
    "users": "учетные записи клиентов и сотрудников",
    "equipment": "карточки позиций спортивного инвентаря",
    "equipmentRates": "тарифы аренды и залог по позициям инвентаря",
    "rentalPointEquipment": "остатки инвентаря по пунктам проката, размерам и состояниям",
    "images": "локальные графические ресурсы приложения",
    "equipmentPhotos": "связь между позициями инвентаря и изображениями",
    "rentOrders": "заголовки заказов аренды",
    "orderItems": "позиции заказов аренды",
    "payments": "платежные записи по заказам",
}

DB_FIELD_DESCRIPTION_DEFAULTS = {
    "id": "Первичный ключ записи.",
    "title": "Текстовое наименование сущности или справочного значения.",
    "code": "Краткий символьный код справочного значения.",
    "unitHours": "Продолжительность одного тарифного интервала в часах.",
    "description": "Текстовое описание записи.",
    "name": "Наименование объекта.",
    "address": "Почтовый адрес объекта.",
    "phone": "Контактный телефон.",
    "idRole": "Ссылка на роль пользователя.",
    "firstName": "Имя пользователя.",
    "lastName": "Фамилия пользователя.",
    "email": "Адрес электронной почты пользователя.",
    "passwordHash": "Строковое представление пароля или его демонстрационного хеша.",
    "dateCreated": "Дата и время создания записи.",
    "idCategory": "Ссылка на категорию инвентаря.",
    "idBrand": "Ссылка на бренд оборудования.",
    "idType": "Ссылка на тип инвентаря.",
    "model": "Модель или артикул позиции инвентаря.",
    "idEquipment": "Ссылка на позицию инвентаря.",
    "idRentalType": "Ссылка на тип тарифа аренды.",
    "price": "Стоимость аренды за выбранный тарифный период.",
    "deposit": "Сумма залога по позиции инвентаря.",
    "idRentalPoint": "Ссылка на пункт проката.",
    "idEquipmentCondition": "Ссылка на состояние инвентаря.",
    "idSize": "Ссылка на размер инвентаря.",
    "totalQuantity": "Общее количество единиц инвентаря в точке проката.",
    "availableQuantity": "Количество единиц, доступных для выдачи в аренду.",
    "url": "Путь или URI к локальному графическому ресурсу.",
    "idImage": "Ссылка на изображение.",
    "idUser": "Ссылка на пользователя, оформившего заказ.",
    "idStatus": "Ссылка на статус записи.",
    "idRentalPointIssue": "Ссылка на пункт выдачи инвентаря.",
    "idRentalPointReturn": "Ссылка на пункт возврата инвентаря.",
    "dateStart": "Дата и время начала аренды.",
    "dateEnd": "Дата и время завершения аренды.",
    "amount": "Денежная сумма по записи.",
    "depositAmount": "Суммарный залог по заказу.",
    "idOrder": "Ссылка на заказ аренды.",
    "idRentalPointEquipment": "Ссылка на конкретный остаток инвентаря в пункте проката.",
    "quantity": "Количество единиц в составе заказа.",
    "pricePerUnit": "Цена одной единицы в составе заказа.",
    "idPaymentMethod": "Ссылка на способ оплаты.",
}

DB_FIELD_DESCRIPTION_OVERRIDES = {
    ("rentalPoints", "name"): "Наименование пункта проката.",
    ("categories", "description"): "Развернутое описание категории инвентаря.",
    ("equipment", "title"): "Наименование позиции инвентаря, отображаемое в каталоге.",
    ("equipment", "description"): "Описание характеристик и особенностей позиции инвентаря.",
    ("users", "phone"): "Контактный телефон пользователя.",
    ("users", "passwordHash"): "Пароль или его строковое представление, используемое в демонстрационной авторизации.",
    ("rentOrders", "idStatus"): "Ссылка на текущий статус заказа аренды.",
    ("payments", "idStatus"): "Ссылка на текущий статус платежа.",
    ("rentOrders", "description"): "Комментарий к заказу аренды.",
    ("rentOrders", "amount"): "Итоговая стоимость аренды без учета залога.",
    ("orderItems", "amount"): "Стоимость конкретной позиции заказа.",
    ("payments", "amount"): "Сумма платежа по заказу.",
    ("images", "url"): "Путь к изображению, встроенному в приложение или поставляемому вместе с ним.",
}

DB_DETAIL_GROUPS = (
    (
        "Справочники и классификаторы базы данных",
        (
            "roles",
            "orderStatuses",
            "paymentMethods",
            "paymentStatuses",
            "equipmentConditions",
            "rentalTypes",
            "categories",
            "brands",
            "equipmentTypes",
            "equipmentSizes",
        ),
    ),
    (
        "Операционные сущности базы данных",
        (
            "rentalPoints",
            "users",
            "equipment",
            "equipmentRates",
            "rentalPointEquipment",
            "images",
            "equipmentPhotos",
            "rentOrders",
            "orderItems",
            "payments",
        ),
    ),
)

DB_DIAGRAM_LAYOUT = {
    "roles": (40, 60),
    "orderStatuses": (40, 300),
    "paymentMethods": (40, 540),
    "paymentStatuses": (40, 780),
    "rentalTypes": (40, 1020),
    "users": (420, 60),
    "rentOrders": (420, 430),
    "payments": (420, 820),
    "rentalPoints": (800, 60),
    "rentalPointEquipment": (800, 430),
    "orderItems": (800, 890),
    "categories": (1180, 60),
    "brands": (1180, 280),
    "equipmentTypes": (1180, 500),
    "equipmentSizes": (1180, 720),
    "equipmentConditions": (1180, 940),
    "equipment": (1560, 60),
    "equipmentRates": (1560, 430),
    "images": (1560, 760),
    "equipmentPhotos": (1560, 980),
}

DB_DIAGRAM_GROUPS = {
    "roles": "reference",
    "orderStatuses": "reference",
    "paymentMethods": "reference",
    "paymentStatuses": "reference",
    "equipmentConditions": "reference",
    "rentalTypes": "reference",
    "categories": "catalog",
    "brands": "catalog",
    "equipmentTypes": "catalog",
    "equipmentSizes": "catalog",
    "rentalPoints": "core",
    "users": "core",
    "equipment": "core",
    "equipmentRates": "operations",
    "rentalPointEquipment": "operations",
    "images": "media",
    "equipmentPhotos": "media",
    "rentOrders": "operations",
    "orderItems": "operations",
    "payments": "operations",
}

DB_DIAGRAM_FIELDS = {
    "roles": ("id", "title"),
    "orderStatuses": ("id", "title"),
    "paymentMethods": ("id", "title"),
    "paymentStatuses": ("id", "title"),
    "equipmentConditions": ("id", "title"),
    "rentalTypes": ("id", "title", "code", "unitHours"),
    "categories": ("id", "title"),
    "brands": ("id", "title"),
    "equipmentTypes": ("id", "title"),
    "equipmentSizes": ("id", "title"),
    "rentalPoints": ("id", "name", "address"),
    "users": ("id", "idRole", "email", "phone"),
    "equipment": ("id", "idCategory", "idBrand", "idType", "title"),
    "equipmentRates": ("id", "idEquipment", "idRentalType", "price", "deposit"),
    "rentalPointEquipment": ("id", "idRentalPoint", "idEquipment", "idSize", "availableQuantity"),
    "images": ("id", "url"),
    "equipmentPhotos": ("id", "idEquipment", "idImage"),
    "rentOrders": ("id", "idUser", "idStatus", "dateStart", "amount"),
    "orderItems": ("id", "idOrder", "idRentalPointEquipment", "quantity", "amount"),
    "payments": ("id", "idOrder", "idPaymentMethod", "idStatus", "amount"),
}


@dataclass(frozen=True)
class DocumentVariant:
    slug: str
    platform_kind: str
    work_label_genitive: str
    subtitle_line: str
    topic: str
    app_term_genitive: str
    app_term_nominative: str
    scenario_focus: str
    ui_note: str
    testing_note: str
    limitation_note: str
    output_docx_name: str
    output_pdf_name: str


VARIANTS = (
    DocumentVariant(
        slug="mobile",
        platform_kind="mobile",
        work_label_genitive="курсовой работы",
        subtitle_line="к курсовой работе",
        topic="Разработка мобильного приложения для аренды спортивного инвентаря",
        app_term_genitive="мобильного приложения",
        app_term_nominative="мобильное приложение",
        scenario_focus=(
            "В данной версии пояснительной записки основной акцент сделан на мобильном сценарии "
            "использования и кроссплатформенной природе решения .NET MAUI."
        ),
        ui_note=(
            "Поскольку проект использует единую XAML-разметку .NET MAUI, в пояснительной записке "
            "используются скриншоты Windows-сборки, отражающие те же экраны и пользовательские "
            "сценарии, что и мобильные цели Android и iOS."
        ),
        testing_note=(
            "Проект ориентирован на мобильные цели net9.0-android и net9.0-ios, однако для "
            "воспроизводимой локальной проверки, подготовки скриншотов и отладки в репозитории "
            "подтверждена сборка Windows-цели net9.0-windows10.0.19041.0. Общая XAML-разметка, "
            "ViewModel и сервисы позволяют рассматривать данную сборку как эквивалентный "
            "демонстрационный контур мобильной версии."
        ),
        limitation_note=(
            "Практическая демонстрация в репозитории выполнена преимущественно на Windows-сборке, "
            "хотя кодовая база ориентирована на мобильные цели Android и iOS."
        ),
        output_docx_name="Пояснительная записка - мобильное приложение SportRent.docx",
        output_pdf_name="Пояснительная записка - мобильное приложение SportRent.pdf",
    ),
    DocumentVariant(
        slug="desktop",
        platform_kind="desktop",
        work_label_genitive="курсовой работы",
        subtitle_line="к курсовой работе",
        topic="Разработка десктопного приложения для аренды спортивного инвентаря",
        app_term_genitive="десктопного приложения",
        app_term_nominative="десктопное приложение",
        scenario_focus=(
            "В данной версии пояснительной записки основной акцент сделан на настольной "
            "Windows-сборке кроссплатформенного проекта SportRent."
        ),
        ui_note=(
            "Для десктопной версии в пояснительной записке используются скриншоты Windows-сборки, "
            "так как именно эта конфигурация была применена для отладки, проверки сценариев и "
            "получения иллюстративного материала."
        ),
        testing_note=(
            "Проверка работоспособности выполнялась на Windows-цели net9.0-windows10.0.19041.0. "
            "Именно эта конфигурация рассматривается как десктопная версия приложения в текущей "
            "пояснительной записке. При этом исходная кодовая база остается общей с мобильными "
            "и другими платформенными целями .NET MAUI."
        ),
        limitation_note="Основной подтвержденный сценарий эксплуатации относится к Windows-сборке MAUI.",
        output_docx_name="Пояснительная записка - десктопное приложение SportRent.docx",
        output_pdf_name="Пояснительная записка - десктопное приложение SportRent.pdf",
    ),
    DocumentVariant(
        slug="mobile-diploma",
        platform_kind="mobile",
        work_label_genitive="дипломного проекта",
        subtitle_line="к дипломному проекту",
        topic="Разработка мобильного приложения для аренды спортивного инвентаря",
        app_term_genitive="мобильного приложения",
        app_term_nominative="мобильное приложение",
        scenario_focus=(
            "В дипломной версии пояснительной записки основной акцент сделан на мобильном "
            "сценарии использования, практической значимости решения и завершенности "
            "пользовательского маршрута аренды в кроссплатформенном приложении .NET MAUI."
        ),
        ui_note=(
            "Поскольку проект использует единую XAML-разметку .NET MAUI, в дипломной "
            "пояснительной записке используются скриншоты Windows-сборки, отражающие те же "
            "экраны и пользовательские сценарии, что и мобильные цели Android и iOS."
        ),
        testing_note=(
            "Проект ориентирован на мобильные цели net9.0-android и net9.0-ios, однако для "
            "воспроизводимой локальной проверки, подготовки скриншотов и отладки в репозитории "
            "подтверждена сборка Windows-цели net9.0-windows10.0.19041.0. Общая XAML-разметка, "
            "ViewModel и сервисы позволяют рассматривать данную сборку как эквивалентный "
            "демонстрационный контур мобильной версии дипломного проекта."
        ),
        limitation_note=(
            "Практическая демонстрация в репозитории выполнена преимущественно на Windows-сборке, "
            "хотя кодовая база ориентирована на мобильные цели Android и iOS."
        ),
        output_docx_name="Пояснительная записка - дипломный проект мобильное приложение SportRent.docx",
        output_pdf_name="Пояснительная записка - дипломный проект мобильное приложение SportRent.pdf",
    ),
)

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
SCREENS_DIR.mkdir(parents=True, exist_ok=True)
TMP_DIR.mkdir(parents=True, exist_ok=True)


def validate_runtime_inputs() -> None:
    missing_inputs: list[str] = []
    if not TEMPLATE_PATH.exists():
        missing_inputs.append(
            f"Word template not found: {TEMPLATE_PATH}. "
            "Set SPORTRENT_NOTE_TEMPLATE_PATH to a valid .dotm file."
        )
    if not DB_PATH.exists():
        missing_inputs.append(
            f"SQLite database not found: {DB_PATH}. "
            "Generate sportRent.db or set SPORTRENT_NOTE_DB_PATH to an existing file."
        )
    if missing_inputs:
        raise FileNotFoundError("\n".join(missing_inputs))


def ensure_base_docx_from_template(template_path: Path, target_path: Path) -> None:
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    word.AutomationSecurity = 3

    document = None
    try:
        document = word.Documents.Add(str(template_path))
        document.SaveAs(str(target_path), FileFormat=16)
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


def export_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    word.AutomationSecurity = 3

    document = None
    try:
        document = word.Documents.Open(str(docx_path))
        document.Fields.Update()
        for toc in document.TablesOfContents:
            toc.Update()
        document.Save()
        document.ExportAsFixedFormat(str(pdf_path), 17)
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


def load_context() -> dict:
    return json.loads(CONTEXT_PATH.read_text(encoding="utf-8"))


def load_facts() -> dict:
    with sqlite3.connect(DB_PATH) as connection:
        cursor = connection.cursor()
        return {
            "table_count": cursor.execute(
                "SELECT count(*) FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'"
            ).fetchone()[0],
            "equipment_count": cursor.execute("SELECT count(*) FROM equipment").fetchone()[0],
            "category_count": cursor.execute("SELECT count(*) FROM categories").fetchone()[0],
            "rental_point_count": cursor.execute("SELECT count(*) FROM rentalPoints").fetchone()[0],
            "user_count": cursor.execute("SELECT count(*) FROM users").fetchone()[0],
            "order_count": cursor.execute("SELECT count(*) FROM rentOrders").fetchone()[0],
            "order_item_count": cursor.execute("SELECT count(*) FROM orderItems").fetchone()[0],
            "payment_count": cursor.execute("SELECT count(*) FROM payments").fetchone()[0],
            "equipment_titles": [
                row[0] for row in cursor.execute("SELECT title FROM equipment ORDER BY id").fetchall()
            ],
        }


def load_db_schema() -> dict[str, dict[str, Any]]:
    schema: dict[str, dict[str, Any]] = {}
    with sqlite3.connect(DB_PATH) as connection:
        cursor = connection.cursor()
        tables = [
            row[0]
            for row in cursor.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name"
            ).fetchall()
        ]
        for table_name in tables:
            columns = [
                {
                    "name": name,
                    "type": data_type or "TEXT",
                    "not_null": bool(not_null or primary_key),
                    "default": default_value,
                    "primary_key": bool(primary_key),
                }
                for _, name, data_type, not_null, default_value, primary_key in cursor.execute(
                    f"PRAGMA table_info('{table_name}')"
                ).fetchall()
            ]
            foreign_keys = [
                {
                    "table": ref_table,
                    "from": from_column,
                    "to": to_column,
                    "on_delete": on_delete,
                }
                for _, _, ref_table, from_column, to_column, _, on_delete, _ in cursor.execute(
                    f"PRAGMA foreign_key_list('{table_name}')"
                ).fetchall()
            ]
            unique_columns: set[str] = set()
            unique_sets: list[tuple[str, ...]] = []
            for _, index_name, is_unique, origin, _ in cursor.execute(
                f"PRAGMA index_list('{table_name}')"
            ).fetchall():
                if not is_unique or origin == "pk":
                    continue
                index_columns = tuple(
                    row[2] for row in cursor.execute(f"PRAGMA index_info('{index_name}')").fetchall()
                )
                if len(index_columns) == 1:
                    unique_columns.add(index_columns[0])
                elif index_columns:
                    unique_sets.append(index_columns)
            schema[table_name] = {
                "columns": columns,
                "foreign_keys": foreign_keys,
                "unique_columns": unique_columns,
                "unique_sets": unique_sets,
            }
    return schema


def get_column_description(table_name: str, column_name: str) -> str:
    override = DB_FIELD_DESCRIPTION_OVERRIDES.get((table_name, column_name))
    if override is not None:
        return override
    return DB_FIELD_DESCRIPTION_DEFAULTS.get(column_name, "Поле хранения служебных или предметных данных.")


def build_column_constraints(schema_table: dict[str, Any], column: dict[str, Any]) -> str:
    constraints: list[str] = []
    if column["primary_key"]:
        constraints.append("PK")
    fk_map = {fk["from"]: fk for fk in schema_table["foreign_keys"]}
    fk = fk_map.get(column["name"])
    if fk is not None:
        fk_text = f"FK -> {fk['table']}.{fk['to']}"
        if fk["on_delete"] == "CASCADE":
            fk_text += ", DELETE CASCADE"
        constraints.append(fk_text)
    if column["name"] in schema_table["unique_columns"]:
        constraints.append("UNIQUE")
    if column["not_null"] and not column["primary_key"]:
        constraints.append("NOT NULL")
    if column["default"] is not None:
        constraints.append(f"DEFAULT {column['default']}")
    return "; ".join(constraints) if constraints else "Без специальных ограничений."


def build_db_field_rows(table_name: str, schema_table: dict[str, Any]) -> list[list[str]]:
    rows = [["Поле", "Тип", "Назначение", "Ограничения и связи"]]
    for column in schema_table["columns"]:
        data_type = "INT" if column["type"].upper() == "INTEGER" else column["type"]
        rows.append(
            [
                column["name"],
                data_type,
                get_column_description(table_name, column["name"]),
                build_column_constraints(schema_table, column),
            ]
        )
    return rows


def format_unique_sets(schema_table: dict[str, Any]) -> str | None:
    unique_sets = schema_table["unique_sets"]
    if not unique_sets:
        return None
    return "; ".join(f"UNIQUE({', '.join(columns)})" for columns in unique_sets)


def get_font(size: int, *, bold: bool = False):
    candidate_names = (
        ("timesbd.ttf", "arialbd.ttf", "seguisb.ttf")
        if bold
        else ("times.ttf", "arial.ttf", "segoeui.ttf")
    )
    for font_name in candidate_names:
        font_path = Path(r"C:\Windows\Fonts") / font_name
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str) -> None:
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        middle = ((start[0] + end[0]) // 2, start[1])
        points = [start, middle, (middle[0], end[1]), end]
    else:
        middle = (start[0], (start[1] + end[1]) // 2)
        points = [start, middle, (end[0], middle[1]), end]
    draw.line(points, fill=color, width=4)

    before_end = points[-2]
    dx = end[0] - before_end[0]
    dy = end[1] - before_end[1]
    if abs(dx) >= abs(dy):
        arrow_points = [
            end,
            (end[0] - 16 if dx > 0 else end[0] + 16, end[1] - 8),
            (end[0] - 16 if dx > 0 else end[0] + 16, end[1] + 8),
        ]
    else:
        arrow_points = [
            end,
            (end[0] - 8, end[1] - 16 if dy > 0 else end[1] + 16),
            (end[0] + 8, end[1] - 16 if dy > 0 else end[1] + 16),
        ]
    draw.polygon(arrow_points, fill=color)


def build_schema_figure(schema: dict[str, dict[str, Any]]) -> Path:
    target_path = TMP_DIR / "db-schema-overview.png"
    image = Image.new("RGB", (1920, 1340), "#fbfbf8")
    draw = ImageDraw.Draw(image)
    title_font = get_font(26, bold=True)
    header_font = get_font(20, bold=True)
    body_font = get_font(17)
    arrow_color = "#7d8792"
    box_width = 320
    header_height = 42
    row_height = 28
    colors = {
        "reference": ("#dfe9f5", "#f7fbff"),
        "catalog": ("#efe9db", "#fcf8ee"),
        "core": ("#e4f0e0", "#f7fcf5"),
        "operations": ("#f7e5d5", "#fff8f1"),
        "media": ("#eadff2", "#fbf8ff"),
    }
    boxes: dict[str, tuple[int, int, int, int]] = {}
    display_lines: dict[str, list[str]] = {}

    draw.text((40, 18), "Логическая схема базы данных SportRent", font=title_font, fill="#222222")

    for table_name, (x, y) in DB_DIAGRAM_LAYOUT.items():
        schema_table = schema[table_name]
        columns_by_name = {column["name"]: column for column in schema_table["columns"]}
        fk_names = {fk["from"] for fk in schema_table["foreign_keys"]}
        lines: list[str] = []
        for field_name in DB_DIAGRAM_FIELDS[table_name]:
            column = columns_by_name.get(field_name)
            if column is None:
                continue
            prefix = "PK " if column["primary_key"] else "FK " if field_name in fk_names else ""
            lines.append(f"{prefix}{field_name}")
        if len(lines) < len(schema_table["columns"]):
            lines.append("...")
        height = header_height + row_height * len(lines) + 16
        boxes[table_name] = (x, y, x + box_width, y + height)
        display_lines[table_name] = lines

    drawn_edges: set[tuple[str, str]] = set()
    for child_table, schema_table in schema.items():
        child_box = boxes.get(child_table)
        if child_box is None:
            continue
        child_center = ((child_box[0] + child_box[2]) // 2, (child_box[1] + child_box[3]) // 2)
        for fk in schema_table["foreign_keys"]:
            parent_table = fk["table"]
            parent_box = boxes.get(parent_table)
            edge = (child_table, parent_table)
            if parent_box is None or edge in drawn_edges:
                continue
            drawn_edges.add(edge)
            parent_center = ((parent_box[0] + parent_box[2]) // 2, (parent_box[1] + parent_box[3]) // 2)
            if parent_center[0] >= child_center[0]:
                start = (child_box[2], child_center[1])
                end = (parent_box[0], parent_center[1])
            else:
                start = (child_box[0], child_center[1])
                end = (parent_box[2], parent_center[1])
            if abs(parent_center[0] - child_center[0]) < 40:
                if parent_center[1] >= child_center[1]:
                    start = (child_center[0], child_box[3])
                    end = (parent_center[0], parent_box[1])
                else:
                    start = (child_center[0], child_box[1])
                    end = (parent_center[0], parent_box[3])
            draw_arrow(draw, start, end, arrow_color)

    for table_name, (x1, y1, x2, y2) in boxes.items():
        group = DB_DIAGRAM_GROUPS[table_name]
        header_color, body_color = colors[group]
        draw.rounded_rectangle((x1, y1, x2, y2), radius=10, fill=body_color, outline="#444444", width=2)
        draw.rounded_rectangle((x1, y1, x2, y1 + header_height), radius=10, fill=header_color, outline="#444444", width=2)
        draw.line((x1, y1 + header_height, x2, y1 + header_height), fill="#444444", width=2)
        draw.text((x1 + 12, y1 + 10), table_name, font=header_font, fill="#1c1c1c")
        current_y = y1 + header_height + 8
        for line in display_lines[table_name]:
            draw.text((x1 + 12, current_y), line, font=body_font, fill="#333333")
            current_y += row_height

    image.save(target_path)
    return target_path


def clear_document(document: Document) -> None:
    body = document.element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def get_style_name(document: Document, preferred: list[str], fallback: str) -> str:
    style_names = {style.name for style in document.styles}
    for name in preferred:
        if name in style_names:
            return name
    return fallback


def get_existing_style(document: Document, preferred: list[str]):
    style_names = {style.name for style in document.styles}
    for name in preferred:
        if name in style_names:
            return document.styles[name]
    return None


def apply_run_font(run, *, bold: bool | None = None, size: int | None = None) -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    if rpr.rFonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    paragraph.add_run()._r.append(begin)
    paragraph.add_run()._r.append(instr)
    paragraph.add_run()._r.append(separate)
    if placeholder:
        placeholder_run = paragraph.add_run(placeholder)
        apply_run_font(placeholder_run, size=12)
    paragraph.add_run()._r.append(end)


def clear_story_container(container) -> None:
    if hasattr(container, "is_linked_to_previous"):
        container.is_linked_to_previous = False
    container_element = container._element
    for child in list(container_element):
        container_element.remove(child)


def clear_section_footers(section, body_style: str) -> None:
    for footer in (section.footer, section.first_page_footer):
        clear_story_container(footer)
        paragraph = footer.add_paragraph()
        paragraph.style = body_style
        paragraph.paragraph_format.first_line_indent = None


def set_default_section(section, *, different_first_page: bool) -> None:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)
    section.different_first_page_header_footer = different_first_page


def ensure_style_font(style, *, font_name: str = "Times New Roman", size: int = 14, bold: bool | None = None) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    if rpr.rFonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rpr.rFonts.set(qn("w:ascii"), font_name)
    rpr.rFonts.set(qn("w:hAnsi"), font_name)
    rpr.rFonts.set(qn("w:eastAsia"), font_name)
    rpr.rFonts.set(qn("w:cs"), font_name)


def normalize_toc_styles(document: Document) -> None:
    toc_heading = get_existing_style(document, ["TOC Heading"])
    if toc_heading is not None:
        ensure_style_font(toc_heading, size=14, bold=False)
        toc_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_heading.paragraph_format.first_line_indent = None
        toc_heading.paragraph_format.line_spacing = 1.5
        toc_heading.paragraph_format.space_before = Pt(0)
        toc_heading.paragraph_format.space_after = Pt(0)

    for level in range(1, 10):
        toc_style = get_existing_style(document, [f"toc {level}", f"TOC {level}"])
        if toc_style is None:
            continue
        ensure_style_font(toc_style, size=14, bold=False)
        toc_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc_style.paragraph_format.first_line_indent = None
        toc_style.paragraph_format.line_spacing = 1.5
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(0)


def prepare_footer(section, body_style: str) -> None:
    footer = section.footer
    clear_story_container(footer)
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.style = body_style
    paragraph.paragraph_format.first_line_indent = None
    add_field(paragraph, "PAGE", "1")
    if section.different_first_page_header_footer:
        first_page_footer = section.first_page_footer
        clear_story_container(first_page_footer)
        paragraph = first_page_footer.add_paragraph()
        paragraph.style = body_style
        paragraph.paragraph_format.first_line_indent = None


def make_paragraph(
    document: Document,
    text: str = "",
    *,
    style: str,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_cm: float | None = 1.25,
    bold: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.alignment = alignment
    if first_line_indent_cm is None:
        paragraph.paragraph_format.first_line_indent = None
    else:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    run = paragraph.add_run(text)
    apply_run_font(run, bold=bold)


def make_blank(document: Document, body_style: str, count: int = 1) -> None:
    for _ in range(count):
        make_paragraph(document, "", style=body_style, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_cm=None)


def set_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for child in list(ppr):
        if child.tag == qn("w:outlineLvl"):
            ppr.remove(child)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    ppr.append(outline)


def make_chapter_title(document: Document, title: str, style: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    set_outline_level(paragraph, 0)
    run = paragraph.add_run(title)
    apply_run_font(run, bold=True, size=16)


def make_toc_title(document: Document, title: str, style: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    set_outline_level(paragraph, 0)
    run = paragraph.add_run(title)
    apply_run_font(run, bold=False, size=14)


def make_section_title(document: Document, title: str, style: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    set_outline_level(paragraph, 1)
    run = paragraph.add_run(title)
    apply_run_font(run, bold=True, size=14)


def make_centered(document: Document, text: str, style: str, *, bold: bool = False, size: int = 14) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run(text)
    apply_run_font(run, bold=bold, size=size)


def make_list_item(document: Document, text: str, body_style: str) -> None:
    make_paragraph(
        document,
        text,
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_cm=None,
    )


def style_table(table, *, font_size: int = 12) -> None:
    try:
        table.style = "Сетка таблицы"
    except KeyError:
        table.style = "Table Grid"

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = None
                for run in paragraph.runs:
                    apply_run_font(run, bold=row_index == 0, size=font_size)


def make_table_with_caption(
    document: Document,
    caption: str,
    rows: list[list[str]],
    caption_style: str,
    body_style: str,
    index: int,
    *,
    font_size: int = 12,
) -> None:
    number_paragraph = document.add_paragraph()
    number_paragraph.style = caption_style
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_paragraph.paragraph_format.first_line_indent = None
    number_run = number_paragraph.add_run(f"Таблица {index}")
    apply_run_font(number_run, size=12)
    make_centered(document, caption, caption_style, size=12)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    default_widths = {
        2: [6.2, 9.0],
        3: [4.0, 7.4, 3.8],
        4: [4.0, 2.6, 4.6, 4.0],
    }
    column_widths = default_widths.get(len(rows[0]))
    if column_widths is not None:
        for column_index, width_cm in enumerate(column_widths):
            table.columns[column_index].width = Cm(width_cm)
            for row in table.rows:
                row.cells[column_index].width = Cm(width_cm)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.style = body_style
                paragraph.paragraph_format.first_line_indent = None
                for run in paragraph.runs:
                    apply_run_font(run, bold=row_index == 0, size=font_size)
    style_table(table, font_size=font_size)
    make_blank(document, body_style)


def make_figure(
    document: Document,
    image_path: Path,
    caption: str,
    body_style: str,
    caption_style: str,
    index: int,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = body_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.2))
    make_centered(document, f"Рис. {index}. {caption}", caption_style, size=12)


def add_db_field_descriptions(
    document: Document,
    schema: dict[str, dict[str, Any]],
    caption_style: str,
    body_style: str,
    start_index: int,
) -> int:
    table_index = start_index
    make_paragraph(
        document,
        "Подробное описание полей приведено ниже. Наименования таблиц и колонок сохранены в том виде, в котором они реализованы в физической модели SQLite и используются прикладным кодом.",
        style=body_style,
    )
    for group_title, table_names in DB_DETAIL_GROUPS:
        make_paragraph(
            document,
            group_title,
            style=body_style,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent_cm=None,
            bold=True,
        )
        for table_name in table_names:
            make_table_with_caption(
                document,
                f"Структура таблицы {table_name}: {DB_TABLE_DESCRIPTION_MAP[table_name]}",
                build_db_field_rows(table_name, schema[table_name]),
                caption_style,
                body_style,
                table_index,
                font_size=10,
            )
            table_index += 1
            unique_sets = format_unique_sets(schema[table_name])
            if unique_sets is not None:
                make_paragraph(
                    document,
                    f"Для таблицы {table_name} дополнительно заданы составные ограничения уникальности: {unique_sets}.",
                    style=body_style,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_indent_cm=None,
                )
    return table_index


def add_toc(document: Document, body_style: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = body_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    add_field(
        paragraph,
        r'TOC \o "1-2" \h \z \u',
        "Оглавление будет обновлено при открытии документа.",
    )


def add_main_section(document: Document, body_style: str) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_default_section(section, different_first_page=False)
    prepare_footer(section, body_style)


def add_appendix_heading(document: Document, body_style: str, appendix_label: str, title: str) -> None:
    appendix_paragraph = document.add_paragraph()
    appendix_paragraph.style = body_style
    appendix_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    appendix_paragraph.paragraph_format.first_line_indent = None
    appendix_run = appendix_paragraph.add_run(f"Приложение {appendix_label}")
    apply_run_font(appendix_run, bold=True)

    make_centered(document, title, body_style, bold=True)
    make_blank(document, body_style)


def make_code_listing(document: Document, source_path: Path, body_style: str) -> None:
    code_text = source_path.read_text(encoding="utf-8").replace("\t", "    ")
    numbered_lines = [
        f"{line_no:04}: {line}" if line else f"{line_no:04}:"
        for line_no, line in enumerate(code_text.splitlines(), start=1)
    ]
    paragraph = document.add_paragraph()
    paragraph.style = body_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run("\n".join(numbered_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    rpr = run._element.get_or_add_rPr()
    if rpr.rFonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rpr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.color.rgb = RGBColor(0, 0, 0)


def verify_screenshots_exist() -> None:
    missing = [path for path in SCREEN_FILES.values() if not path.exists()]
    if missing:
        raise FileNotFoundError(f"Missing screenshots: {missing}")


def get_intro_paragraphs(context: dict, variant: DocumentVariant) -> list[str]:
    return [
        (
            "Актуальность разработки приложений для аренды спортивного инвентаря определяется "
            "ростом спроса на краткосрочное пользование оборудованием без необходимости его покупки. "
            "Для клиента важно быстро найти нужную позицию, проверить доступность в конкретном "
            "пункте проката и получить понятные условия аренды."
        ),
        (
            f"Тема {variant.work_label_genitive} связана с созданием {variant.app_term_genitive} на базе "
            "кроссплатформенной технологии .NET MAUI, объединяющего проектирование базы данных "
            "и реализацию пользовательского интерфейса. В рамках проекта SportRent реализован "
            "автономный демонстрационный сценарий, в котором основные действия выполняются на "
            "базе локальной SQLite без серверной части."
        ),
        (
            f"Целью работы является разработка {variant.app_term_genitive} SportRent, "
            "обеспечивающего авторизацию пользователя, просмотр каталога, получение сведений об "
            "инвентаре, оформление аренды, просмотр истории заказов и личного кабинета."
        ),
        (
            "Работа выполнена на примере учебного сервиса проката спортивного инвентаря SportRent, "
            "моделирующего деятельность пункта проката с локальной базой данных и типовыми "
            "пользовательскими сценариями."
        ),
        variant.scenario_focus,
        "Для достижения поставленной цели были решены следующие задачи:",
    ]


def get_task_items(variant: DocumentVariant) -> list[str]:
    task_four = (
        "4. Создано кроссплатформенное мобильное приложение на платформе .NET MAUI."
        if variant.platform_kind == "mobile"
        else "4. Создана десктопная сборка приложения на платформе .NET MAUI для Windows."
    )
    return [
        "1. Выполнен анализ предметной области аренды спортивного инвентаря.",
        "2. Спроектирована нормализованная структура реляционной базы данных.",
        "3. Реализован инструмент автоматической сборки SQLite из SQL-скриптов.",
        task_four,
        "5. Проверена работоспособность ключевых пользовательских сценариев.",
        "6. Выполнен учебный расчет экономической эффективности проекта.",
    ]


def get_requirements_paragraphs(variant: DocumentVariant) -> list[str]:
    platform_requirement = (
        "Для мобильной версии дополнительно важны компактная компоновка элементов, вертикальная "
        "прокрутка контента и удобство сенсорного взаимодействия на экранах ограниченного размера."
        if variant.platform_kind == "mobile"
        else "Для десктопной версии дополнительно важны работа в оконном режиме, рациональное "
        "использование расширенной рабочей области и удобство взаимодействия с помощью клавиатуры и мыши."
    )
    return [
        (
            "На основании анализа предметной области были сформулированы требования к программному "
            "средству. Пользовательская часть должна обеспечивать быстрый вход в систему, поиск и "
            "фильтрацию каталога, просмотр детальной карточки, оформление аренды и получение "
            "информации о ранее созданных заказах."
        ),
        (
            "Система должна работать в автономном режиме, поскольку проект ориентирован на учебную "
            "демонстрацию без зависимости от удаленной инфраструктуры. Поэтому база данных, "
            "изображения инвентаря и логика доступа к данным размещены локально в составе клиента."
        ),
        (
            "Дополнительным требованием стало обеспечение воспроизводимости данных. Для этого в "
            "решении выделен отдельный инструмент сборки базы данных из канонических SQL-скриптов "
            "и подготовлены демонстрационные учетные записи разных ролей."
        ),
        platform_requirement,
    ]


def get_architecture_paragraphs(variant: DocumentVariant) -> list[str]:
    focus_text = (
        "В рамках данной версии записки основной акцент сделан на мобильных пользовательских "
        "сценариях и переносимости экранов между Android и iOS."
        if variant.platform_kind == "mobile"
        else "В рамках данной версии записки основной акцент сделан на Windows-сборке, которая "
        "рассматривается как десктопный вариант эксплуатации приложения."
    )
    return [
        (
            "Решение состоит из трех логически разделенных проектов. Такое разделение позволяет "
            "обособить задачи построения базы данных, работы со scaffolded сущностями и разработки "
            "пользовательского интерфейса."
        ),
        (
            "В клиентском приложении используется упрощенный вариант архитектурного подхода MVVM: "
            "страницы отвечают за визуальную часть, ViewModel управляют состоянием экрана, а сервисы "
            "инкапсулируют доступ к данным и прикладную логику. Такой подход уменьшает связанность "
            "кода и делает сценарии загрузки каталога, авторизации и заказов более прозрачными."
        ),
        (
            "Файл SportRent.Mobile.csproj подтверждает кроссплатформенный характер решения: "
            "проект ориентирован на Android, iOS, Mac Catalyst и Windows. Общая XAML-разметка и "
            "единый слой сервисов обеспечивают использование одной кодовой базы для разных целей сборки."
        ),
        focus_text,
        (
            "Точкой входа служит класс App, который определяет стартовый экран по состоянию "
            "пользовательской сессии. При отсутствии авторизации отображается LoginPage, а после "
            "входа открывается AppShell с вкладками «Каталог», «Заказы» и «Профиль»."
        ),
    ]


def get_interface_intro_paragraphs(variant: DocumentVariant) -> list[str]:
    layout_note = (
        "Интерфейс мобильной версии ориентирован на вертикальный сценарий просмотра: блоки "
        "последовательно выстраиваются сверху вниз, а ключевые действия размещены в видимой зоне "
        "экрана, что упрощает сенсорное взаимодействие."
        if variant.platform_kind == "mobile"
        else "Интерфейс десктопной версии работает в оконном режиме Windows и использует увеличенную "
        "ширину рабочей области для более комфортного просмотра карточек, списков и сводных показателей."
    )
    return [
        (
            "Интерфейс приложения оформлен в едином визуальном стиле и ориентирован на быстрое "
            "прохождение основных сценариев. На экранах используются крупные карточки, сводные "
            "метрики, фильтрация по категориям и визуально выделенные области для стоимости, "
            "залога и доступности."
        ),
        layout_note,
        variant.ui_note,
    ]


def get_testing_paragraphs(variant: DocumentVariant, facts: dict) -> list[str]:
    return [
        variant.testing_note,
        (
            f"В результате проверки установлено, что приложение корректно отображает "
            f"{facts['equipment_count']} позиций каталога, сгруппированных по "
            f"{facts['category_count']} категориям и доступных в {facts['rental_point_count']} "
            "пунктах проката. История заказов показывает подготовленный тестовый заказ, а профиль "
            "пользователя агрегирует статистику по заказам и платежам."
        ),
        (
            "Основной проверочный сценарий включает вход под демонстрационным пользователем, "
            "просмотр каталога, открытие карточки инвентаря, переход к истории заказов и просмотру "
            "личного кабинета. Этот набор действий показывает связность интерфейса, корректность "
            "навигации и согласованность локальных данных."
        ),
    ]


def get_limitations(variant: DocumentVariant) -> list[str]:
    return [
        "1. Авторизация демонстрационная, пароль хранится в seed-данных без криптографического хеширования.",
        "2. Серверная часть и синхронизация с удаленной базой данных отсутствуют.",
        "3. Регистрация новых пользователей и полноценная ролевая модель интерфейса не реализованы.",
        "4. Реальная интеграция с платежной системой отсутствует.",
        "5. Автоматизированные тесты в текущей версии не разработаны.",
        f"6. {variant.limitation_note}",
    ]


def get_economic_cost_rows() -> list[list[str]]:
    return [
        ["Статья затрат", "Основание расчета", "Сумма, руб."],
        ["Анализ и проектирование", "48 ч x 350 руб./ч", "16 800"],
        ["Разработка базы данных и сервисов", "92 ч x 350 руб./ч", "32 200"],
        ["Разработка пользовательского интерфейса", "76 ч x 350 руб./ч", "26 600"],
        ["Тестирование и отладка", "40 ч x 320 руб./ч", "12 800"],
        ["Электроэнергия и амортизация ПК", "Условный расчет", "4 600"],
        ["Итого", "", "93 000"],
    ]


def get_economic_effect_rows() -> list[list[str]]:
    return [
        ["Показатель", "Значение"],
        ["Среднее сокращение времени на оформление одной аренды", "8 минут"],
        ["Среднее количество заявок в месяц", "180"],
        ["Экономия рабочего времени в месяц", "24 часа"],
        ["Средняя стоимость часа работы сотрудника", "350 руб."],
        ["Условная ежемесячная экономия", "8 400 руб."],
        ["Дополнительный эффект от снижения ошибок и повторных операций", "4 000 руб."],
        ["Совокупный экономический эффект в месяц", "12 400 руб."],
        ["Ориентировочный срок окупаемости", "около 7,5 месяцев"],
    ]


def get_conclusion_paragraphs(variant: DocumentVariant) -> list[str]:
    first = (
        f"В ходе выполнения {variant.work_label_genitive} разработано мобильное кроссплатформенное приложение "
        "SportRent для аренды спортивного инвентаря с локальной SQLite-базой данных."
        if variant.platform_kind == "mobile"
        else f"В ходе выполнения {variant.work_label_genitive} разработана десктопная версия кроссплатформенного "
        "приложения SportRent для аренды спортивного инвентаря с локальной SQLite-базой данных."
    )
    return [
        first,
        (
            f"В результате работы спроектирована нормализованная база данных, создан инструмент ее "
            "автоматической сборки из SQL-скриптов, выполнено scaffold-моделирование слоя данных и "
            f"реализовано {variant.app_term_nominative} на .NET MAUI. Пользователю доступны "
            "авторизация, каталог, карточка инвентаря, оформление аренды, история заказов и профиль."
        ),
        (
            f"Поставленная цель достигнута, так как тема {variant.work_label_genitive} «{variant.topic}» "
            "подтверждена фактической реализацией программного средства. Полученный результат может "
            "служить основой для дальнейшего развития проекта, включая добавление серверной части, "
            "улучшение безопасности, расширение ролей пользователей и интеграцию с внешними сервисами."
        ),
    ]


def build_document(variant: DocumentVariant) -> tuple[Path, Path]:
    context = load_context()
    facts = load_facts()
    db_schema = load_db_schema()
    db_schema_figure = build_schema_figure(db_schema)
    verify_screenshots_exist()

    base_docx_path = TMP_DIR / f"{variant.slug}-template-base.docx"
    docx_path = OUTPUT_DIR / variant.output_docx_name
    pdf_path = OUTPUT_DIR / variant.output_pdf_name

    ensure_base_docx_from_template(TEMPLATE_PATH, base_docx_path)
    document = Document(str(base_docx_path))
    clear_document(document)

    body_style = get_style_name(document, ["ГОСТ Текст", "Основной текст", "Normal"], "Normal")
    chapter_style = body_style
    section_style = body_style
    caption_style = get_style_name(document, ["caption", "Caption", "ГОСТ Текст"], body_style)
    normalize_toc_styles(document)

    for section in document.sections:
        set_default_section(section, different_first_page=True)
        clear_section_footers(section, body_style)

    year = str(dt.date.today().year)
    figure_index = 1
    table_index = 1

    make_centered(document, "[Наименование образовательной организации]", body_style, bold=False)
    make_centered(document, "[Кафедра / отделение]", body_style)
    make_blank(document, body_style, 4)
    make_centered(document, "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", body_style, bold=True, size=18)
    make_centered(document, variant.subtitle_line, body_style)
    make_centered(document, "на тему", body_style)
    make_centered(document, f"«{variant.topic}»", body_style, bold=True)
    make_blank(document, body_style, 3)
    make_paragraph(
        document,
        "Выполнил: ____________________",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent_cm=None,
    )
    make_paragraph(
        document,
        "Группа: ____________________",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent_cm=None,
    )
    make_paragraph(
        document,
        "Руководитель: ____________________",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent_cm=None,
    )
    make_paragraph(
        document,
        "Город: ____________________",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent_cm=None,
    )
    make_blank(document, body_style, 3)
    make_centered(document, "[Год]", body_style)

    document.add_page_break()
    make_toc_title(document, "СОДЕРЖАНИЕ", chapter_style)
    add_toc(document, body_style)

    add_main_section(document, body_style)
    make_chapter_title(document, "ВВЕДЕНИЕ", chapter_style)
    for paragraph in get_intro_paragraphs(context, variant):
        make_paragraph(document, paragraph, style=body_style)
    for item in get_task_items(variant):
        make_list_item(document, item, body_style)
    make_paragraph(
        document,
        (
            "Объектом исследования является процесс аренды спортивного инвентаря, а предметом "
            "исследования выступают методы проектирования локальной базы данных и реализации "
            f"{variant.app_term_genitive} для обслуживания данного процесса."
        ),
        style=body_style,
    )
    make_paragraph(
        document,
        "В работе использованы методы анализа предметной области, сравнения аналогичных решений, проектирования структуры данных, программной реализации, тестирования и условного экономического расчета.",
        style=body_style,
    )

    document.add_page_break()
    make_chapter_title(document, "ГЛАВА 1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ ИЗУЧАЕМОЙ ПРОБЛЕМЫ", chapter_style)
    make_section_title(document, "1.1 Характеристика предметной области", section_style)
    chapter_one = [
        "Предметная область включает пользователей, пункты проката, категории и конкретные позиции инвентаря, тарифы аренды, остатки, заказы и платежи. Для одного и того же вида инвентаря необходимо хранить сведения о доступности в разных пунктах, размерах и состояниях, а также учитывать залог и срок аренды.",
        f"В текущей версии проекта в локальной базе данных хранится {facts['category_count']} категорий, {facts['equipment_count']} позиций инвентаря и {facts['rental_point_count']} пункта проката. Пользователю доступны такие примеры оборудования, как {', '.join(facts['equipment_titles'][:5])}. Эти данные достаточны для демонстрации типового потока работы приложения на защите.",
        "Ключевая особенность рассматриваемой предметной области состоит в необходимости связать каталог с реальными остатками инвентаря. Пользователь должен видеть не только карточку товара, но и доступность выбранной позиции в конкретном пункте проката, а также применимые тарифы и размер залога.",
    ]
    for paragraph in chapter_one:
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "1.2 Требования к цифровому сервису аренды", section_style)
    for paragraph in get_requirements_paragraphs(variant):
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "1.3 Обоснование выбора средств разработки", section_style)
    for paragraph in [
        f"Для реализации проекта выбран стек, включающий {context['technology_stack']['platform']} в качестве платформы, {context['technology_stack']['ui']} для построения интерфейса, язык {context['technology_stack']['language']}, разметку {context['technology_stack']['markup']}, СУБД {context['technology_stack']['database']} и Entity Framework Core 9 для слоя данных. Такой набор средств соответствует теме {variant.work_label_genitive} и позволяет объединить разработку интерфейса, локального хранения и прикладной логики в одной технологии.",
        "Выбор SQLite обусловлен малым объемом демонстрационных данных, простотой развертывания и возможностью хранить файл базы непосредственно в составе приложения. .NET MAUI обеспечивает кроссплатформенную основу и единый стек для реализации страниц, навигации и взаимодействия с локальными ресурсами.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Использованные технологии и их назначение",
        [
            ["Технология", "Назначение"],
            [".NET 9", "Платформа разработки и выполнения проекта"],
            [".NET MAUI", "Построение кроссплатформенного пользовательского интерфейса"],
            ["C#", "Реализация прикладной логики и сервисов"],
            ["XAML", "Описание экранов приложения"],
            ["SQLite", "Локальное хранение данных"],
            ["Entity Framework Core 9", "Scaffolded слой данных и модели"],
            ["PowerShell", "Автоматизация сборки базы данных и scaffold-процесса"],
        ],
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1

    document.add_page_break()
    make_chapter_title(document, "ГЛАВА 2. ПРОЕКТИРОВАНИЕ СТРУКТУРЫ ПРИЛОЖЕНИЯ", chapter_style)
    make_section_title(document, "2.1 Проектирование базы данных", section_style)
    for paragraph in [
        f"База данных проекта сформирована как нормализованная реляционная структура и в текущем состоянии содержит {facts['table_count']} таблиц. При проектировании отдельно выделены справочники ролей, категорий, брендов, типов аренды, статусов заказов и платежей, что исключает дублирование данных и упрощает дальнейшее расширение системы.",
        "Основной сущностью каталога является таблица equipment, содержащая сведения о конкретной позиции инвентаря. Тарифы вынесены в отдельную таблицу equipmentRates, что позволяет задавать несколько вариантов аренды для одной позиции. Остатки хранятся в rentalPointEquipment и связаны одновременно с пунктом проката, размером и состоянием оборудования.",
        "Для корректного хранения медиафайлов реализована нормализация изображений через таблицы images и equipmentPhotos. Позиции заказа связаны не просто с инвентарем вообще, а с конкретным остатком из rentalPointEquipment, благодаря чему приложение может корректно уменьшать доступное количество при создании заказа.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Ключевые группы таблиц базы данных",
        [
            ["Группа", "Содержание"],
            ["Пользователи и роли", "roles, users"],
            ["Каталог и классификаторы", "categories, brands, equipmentTypes, equipmentSizes, equipmentConditions"],
            ["Прокат и доступность", "rentalTypes, rentalPoints, equipment, equipmentRates, rentalPointEquipment"],
            ["Изображения", "images, equipmentPhotos"],
            ["Заказы и статусы", "orderStatuses, rentOrders, orderItems"],
            ["Платежи", "paymentMethods, paymentStatuses, payments"],
        ],
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    make_paragraph(
        document,
        "Важным проектным решением стало хранение денежных значений в целочисленном формате, то есть в копейках. Это исключает ошибки округления и делает операции расчета аренды и залога более предсказуемыми.",
        style=body_style,
    )
    make_paragraph(
        document,
        f"На рисунке {figure_index} представлена логическая схема базы данных SportRent. На ней показаны основные таблицы и внешние ключи между пользователями, каталогом, остатками, заказами и платежами.",
        style=body_style,
    )
    make_figure(
        document,
        db_schema_figure,
        "Логическая схема базы данных SportRent",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Схема показывает, что таблица equipment образует центр каталожного контура, rentalPointEquipment связывает каталог с фактическими остатками по точкам проката, а rentOrders, orderItems и payments формируют транзакционный контур аренды. Таблица rentOrders дополнительно содержит две ссылки на rentalPoints, что позволяет отдельно хранить пункт выдачи и пункт возврата инвентаря.",
        style=body_style,
    )
    table_index = add_db_field_descriptions(document, db_schema, caption_style, body_style, table_index)
    make_paragraph(
        document,
        "Представленная детализация фиксирует не только перечень сущностей, но и состав их полей, типы данных, обязательность заполнения и внешние ключи. Это делает описание базы данных в записке согласованным с фактической SQLite-схемой проекта и удобным для последующей доработки приложения.",
        style=body_style,
    )
    make_section_title(document, "2.2 Проектирование архитектуры решения", section_style)
    for paragraph in get_architecture_paragraphs(variant):
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Состав решения SportRent",
        [
            ["Проект", "Назначение"],
            ["SportRent.DbTool", "Сборка SQLite-базы из упорядоченных SQL-скриптов"],
            ["SportRent.Data", "Scaffolded EF Core сущности и DbContext"],
            ["SportRent.Mobile", "Кроссплатформенное приложение с экранами, сервисами и навигацией для мобильной и десктопной сборки"],
        ],
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    make_table_with_caption(
        document,
        "Целевые платформы единой кодовой базы MAUI",
        PLATFORM_ROWS,
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    document.add_page_break()
    make_chapter_title(document, "ГЛАВА 3. РЕАЛИЗАЦИЯ И ТЕСТИРОВАНИЕ ПРИЛОЖЕНИЯ", chapter_style)
    make_section_title(document, "3.1 Реализация основных пользовательских сценариев", section_style)
    for paragraph in [
        "Сервис авторизации обращается к локальной базе данных, загружает демонстрационные учетные записи и выполняет проверку логина и пароля по таблице users с учетом роли пользователя. В учебной версии проекта учетные данные используются в демонстрационном виде, что упрощает проверку интерфейса и базы данных.",
        "Сервис каталога формирует снимок данных для главного экрана, включающий список категорий, оборудование и сводные показатели по количеству позиций, пунктов проката и категорий. Для карточки инвентаря сервис дополнительно извлекает тарифы аренды, величину залога и сведения о доступности по точкам проката.",
        "Наиболее важным с точки зрения прикладной логики является сценарий оформления аренды. В сервисе заказов реализована транзакционная последовательность действий: проверка остатка, уменьшение доступного количества, создание записи заказа, добавление позиции заказа и автоматическое создание платежа со статусом ожидания оплаты. Такой порядок обеспечивает согласованность локальных данных.",
        f"Подготовленные демонстрационные данные содержат {facts['user_count']} учетные записи, {facts['order_count']} тестовый заказ, {facts['order_item_count']} позицию(й) заказа и {facts['payment_count']} платежную запись. Благодаря этому пользователь может показать на защите не только каталог, но и уже заполненную историю аренды.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "3.2 Разработка интерфейса приложения", section_style)
    for paragraph in get_interface_intro_paragraphs(variant):
        make_paragraph(document, paragraph, style=body_style)
    make_paragraph(
        document,
        "Экран авторизации.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "Первым для пользователя открывается экран входа, на котором размещены предзаполненные поля почты и пароля, а также список демонстрационных учетных записей. Такое решение уменьшает время подготовки к защите и позволяет быстро перейти к показу основных функций приложения.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["login"],
        "Экран входа в систему с демонстрационными учетными записями",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Главный экран каталога.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "После успешной авторизации открывается вкладка каталога. На ней отображаются агрегированные показатели по количеству позиций, пунктов проката и категорий, строка поиска, горизонтальный набор фильтров и карточки инвентаря с кратким описанием, стартовой ценой и залогом.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["catalog"],
        "Главный экран каталога спортивного инвентаря",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Карточка инвентаря и переход к оформлению аренды.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "При выборе позиции из каталога пользователь переходит на экран подробной карточки. Здесь собраны сведения о категории, бренде и модели, минимальной стоимости аренды, размере залога, доступности по пунктам проката и доступных тарифах. С этого экрана выполняется переход к отдельной форме оформления аренды, где пользователь выбирает тариф, пункт выдачи, дату начала, количество единиц и сразу видит итоговый расчет платежа.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["details"],
        "Карточка выбранной позиции инвентаря",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Экран истории заказов.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "Раздел заказов показывает уже оформленные аренды и предназначен для проверки результатов транзакционного сценария. На экране видны статус заказа, сроки аренды, сумма аренды, размер залога, итоговый платеж и состояние оплаты, что позволяет наглядно проверить корректность записи данных в локальную базу.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["orders"],
        "Экран истории заказов пользователя",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        "Экран профиля пользователя.",
        style=body_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None,
        bold=True,
    )
    make_paragraph(
        document,
        "Вкладка профиля завершает пользовательский маршрут и содержит персональные данные, роль, контакты, дату регистрации и сводные показатели по заказам и оплатам. Наличие кнопки выхода из аккаунта подтверждает завершенность базового сценария работы с пользовательской сессией.",
        style=body_style,
    )
    make_figure(
        document,
        SCREEN_FILES["profile"],
        "Экран профиля пользователя",
        body_style,
        caption_style,
        figure_index,
    )
    figure_index += 1
    make_paragraph(
        document,
        f"Представленные экраны показывают, что тема {variant.work_label_genitive} соответствует фактически реализованному приложению: пользователь может войти в систему, выбрать инвентарь, ознакомиться с условиями аренды, посмотреть историю заказов и получить сведения о своем профиле.",
        style=body_style,
    )

    document.add_page_break()
    make_section_title(document, "3.3 Тестирование и проверка работоспособности", section_style)
    make_section_title(document, "3.3.1 Сборка локальной базы данных", section_style)
    for paragraph in [
        "Для подготовки базы данных реализован отдельный консольный инструмент SportRent.DbTool. Он выполняет поиск SQL-файлов формата 000_name.sql в каталоге database, сортирует их по имени, создает временный файл SQLite и последовательно выполняет все найденные сценарии в рамках транзакции.",
        "В проекте используются три канонических сценария: 001_recreate.sql, 002_schema.sql и 003_seed.sql. Первый сценарий отвечает за пересоздание структуры, второй - за формирование схемы, третий - за наполнение демонстрационными данными.",
        "После успешного выполнения SQL-скриптов временный файл базы данных заменяет итоговый файл sportRent.db в ресурсах мобильного приложения. Данный подход снижает риск повреждения рабочей базы и делает процесс подготовки данных воспроизводимым.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "3.3.2 Проверка пользовательских сценариев", section_style)
    for paragraph in get_testing_paragraphs(variant, facts):
        make_paragraph(document, paragraph, style=body_style)
    make_section_title(document, "3.3.3 Ограничения текущей версии", section_style)
    make_paragraph(
        document,
        "Рассматриваемая реализация является учебной и не претендует на уровень промышленной системы. Это следует учитывать при оценке результатов и определении направлений дальнейшего развития проекта.",
        style=body_style,
    )
    for item in get_limitations(variant):
        make_list_item(document, item, body_style)
    make_paragraph(
        document,
        "Несмотря на перечисленные ограничения, проект решает поставленные учебные задачи и демонстрирует полный маршрут обработки данных от структуры базы до экранов конечного пользователя.",
        style=body_style,
    )

    document.add_page_break()
    make_chapter_title(document, "ГЛАВА 4. РАСЧЕТ ЭКОНОМИЧЕСКОЙ ЭФФЕКТИВНОСТИ ПРОЕКТА", chapter_style)
    make_section_title(document, "4.1 Анализ затрат на разработку", section_style)
    for paragraph in [
        "Расчет экономической эффективности выполнен в учебных целях и носит оценочный характер. Он позволяет показать, какие трудозатраты и сопутствующие расходы могут сопровождать разработку подобного программного продукта.",
        "В расчет включены основные этапы проектирования, программной реализации, тестирования, а также условные затраты на использование вычислительной техники. Оценка построена на трудоемкости работ, наблюдаемой в ходе разработки учебного проекта SportRent.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Оценка затрат на разработку проекта",
        get_economic_cost_rows(),
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    make_section_title(document, "4.2 Оценка экономической эффективности", section_style)
    for paragraph in [
        "Для оценки эффекта рассмотрен типовой пункт проката, в котором часть операций менеджера переносится в приложение: просмотр каталога, уточнение доступности, расчет стоимости аренды и фиксация заказа. Это уменьшает время обслуживания одного клиента и снижает вероятность повторного ввода данных.",
        "Показатели экономического эффекта основаны на условных, но реалистичных допущениях по количеству обращений и стоимости рабочего времени сотрудника. Такой подход соответствует учебному характеру проекта и позволяет определить ориентировочный срок окупаемости.",
    ]:
        make_paragraph(document, paragraph, style=body_style)
    make_table_with_caption(
        document,
        "Оценка ежемесячного экономического эффекта",
        get_economic_effect_rows(),
        caption_style,
        body_style,
        table_index,
    )
    table_index += 1
    make_paragraph(
        document,
        "Полученные оценки показывают, что даже при учебном масштабе проекта приложение способно сократить операционные затраты пункта проката и окупить условные затраты на разработку в разумный срок. Следовательно, разработка SportRent является не только технически реализуемой, но и экономически целесообразной в модели демонстрационного внедрения.",
        style=body_style,
    )

    document.add_page_break()
    make_chapter_title(document, "ЗАКЛЮЧЕНИЕ", chapter_style)
    for paragraph in get_conclusion_paragraphs(variant):
        make_paragraph(document, paragraph, style=body_style)

    document.add_page_break()
    make_chapter_title(document, "СПИСОК ЛИТЕРАТУРЫ", chapter_style)
    for item in [
        "1. Методические указания к дипломному проектированию. Локальный документ «МУ к ДП.pdf».",
        "2. Репозиторий SportRent. Файл README.md.",
        "3. Репозиторий SportRent. Файл docs/explanatory-note-agent-context.json.",
        "4. Репозиторий SportRent. Файл SportRent.Mobile/SportRent.Mobile.csproj.",
        "5. Репозиторий SportRent. Исходный код проекта SportRent.DbTool, файл Program.cs.",
        "6. Репозиторий SportRent. Скрипт SportRent.Data/scaffold.ps1.",
        "7. Репозиторий SportRent. Исходный код сервиса авторизации SportRent.Mobile/Services/AuthenticationService.cs.",
        "8. Репозиторий SportRent. Исходный код каталожного сервиса SportRent.Mobile/Services/SportRentCatalogService.cs.",
        "9. Репозиторий SportRent. Исходный код сервиса заказов SportRent.Mobile/Services/OrdersService.cs.",
        "10. Репозиторий SportRent. SQL-сценарий структуры базы данных database/002_schema.sql.",
        "11. Репозиторий SportRent. SQL-сценарий начального заполнения database/003_seed.sql.",
        "12. Microsoft Learn. Documentation for .NET MAUI.",
        "13. Microsoft Learn. Documentation for Entity Framework Core.",
        "14. SQLite Documentation.",
    ]:
        make_list_item(document, item, body_style)

    for appendix_label, (appendix_title, appendix_path) in zip(APPENDIX_LABELS, CODE_LISTINGS, strict=True):
        document.add_page_break()
        add_appendix_heading(
            document,
            body_style,
            appendix_label,
            f"{appendix_title} ({appendix_path.name})",
        )
        make_code_listing(document, appendix_path, body_style)

    document.save(docx_path)
    export_to_pdf(docx_path, pdf_path)
    return docx_path, pdf_path


if __name__ == "__main__":
    validate_runtime_inputs()
    for current_variant in VARIANTS:
        current_docx, current_pdf = build_document(current_variant)
        print(f"{current_variant.slug.upper()} DOCX: {current_docx}")
        print(f"{current_variant.slug.upper()} PDF: {current_pdf}")
